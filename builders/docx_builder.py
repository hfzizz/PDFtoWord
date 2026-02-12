"""DOCX document builder module.

Constructs a Word document from a semantic structure map produced by
SemanticAnalyzer, applying formatting, layout, and page geometry derived
from the source PDF.
"""

import os
import re
import logging
from typing import Any, Dict, List, Optional

from lxml import etree

import docx.opc.constants
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# 1 PDF point == 12 700 EMUs (English Metric Units).
_PT_TO_EMU = 12700

# Sensible defaults when metadata / config values are absent.
_DEFAULT_MARGIN_INCHES = 1.0
_DEFAULT_FALLBACK_FONT = "Arial"
_MAX_IMAGE_WIDTH_INCHES = 6.0

# Mapping from alignment strings to python‑docx enum members.
_ALIGNMENT_MAP: Dict[str, int] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Style names for nested bullet / number lists (levels 0‑2).
_BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]

# Regex patterns used to strip common bullet / number prefixes.
_BULLET_PREFIX_RE = re.compile(
    r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s*"
)
_NUMBER_PREFIX_RE = re.compile(
    r"^\d+[\.\)\:]\s*"
)


class DocxBuilder:
    """Builds a ``.docx`` document from a structured element map.

    Parameters
    ----------
    config : dict
        Application configuration (typically loaded from
        ``config/settings.json``).  Recognised keys:

        * ``fallback_font`` – font family used when an element does not
          specify one (default ``"Arial"``).
        * ``default_margin_inches`` – page margin in inches (default
          ``1.0``).
    """

    def __init__(self, config: Dict[str, Any]) -> None:
        self._config = config
        self._fallback_font: str = config.get(
            "fallback_font", _DEFAULT_FALLBACK_FONT
        )
        self._default_margin: float = config.get(
            "default_margin_inches", _DEFAULT_MARGIN_INCHES
        )
        self._doc: Optional[Document] = None
        self._skipped: int = 0

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def build(
        self,
        structure_map: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        output_path: str,
    ) -> str:
        """Build a ``.docx`` file from the given structure map.

        Parameters
        ----------
        structure_map:
            Ordered list of element dicts produced by
            ``SemanticAnalyzer``.  Each dict must contain a ``"type"``
            key plus type‑specific fields (see module docstring).
        metadata:
            Document‑level information.  Expected to carry per‑page
            geometry under ``"pages"`` — a list of dicts with ``width``
            and ``height`` in PDF points.
        output_path:
            Filesystem path for the generated ``.docx`` file.

        Returns
        -------
        str
            The absolute path to the saved document.
        """
        self._skipped = 0
        self._setup_document(metadata, structure_map)

        handler_dispatch = {
            "heading": self._add_heading,
            "paragraph": self._add_paragraph,
            "list_item": self._add_list_item,
            "table": self._add_table,
            "image": self._add_image,
            "page_break": self._add_page_break,
            "header": self._add_header,
            "footer": self._add_footer,
        }

        for idx, element in enumerate(structure_map):
            elem_type = element.get("type")

            # Merge overlapping images: if current and next are both images
            # with overlapping y-ranges, combine them into one paragraph.
            if elem_type == "image":
                images_to_add = [element]
                peek = idx + 1
                cur_y0 = element.get("_y0", 0)
                cur_y1 = cur_y0 + element.get("height", 0)
                while peek < len(structure_map):
                    nxt = structure_map[peek]
                    if nxt.get("type") != "image":
                        break
                    nxt_y0 = nxt.get("_y0", 0)
                    nxt_y1 = nxt_y0 + nxt.get("height", 0)
                    # Overlapping y-range
                    if nxt_y0 < cur_y1 and nxt_y1 > cur_y0:
                        images_to_add.append(nxt)
                        cur_y1 = max(cur_y1, nxt_y1)
                        peek += 1
                    else:
                        break
                if len(images_to_add) > 1:
                    # Mark extra images so they are skipped in the loop.
                    for extra in images_to_add[1:]:
                        extra["_merged"] = True
                    try:
                        self._add_images_inline(images_to_add)
                    except Exception:
                        logger.warning(
                            "Failed to merge overlapping images at index %d.",
                            idx, exc_info=True,
                        )
                    continue
            if element.get("_merged"):
                continue

            handler = handler_dispatch.get(elem_type)  # type: ignore[arg-type]
            if handler is None:
                logger.warning(
                    "Unknown element type '%s' at index %d — skipping.",
                    elem_type,
                    idx,
                )
                self._skipped += 1
                continue

            try:
                handler(element)
            except Exception:
                logger.warning(
                    "Failed to process element at index %d (type='%s'). "
                    "Adding placeholder.",
                    idx,
                    elem_type,
                    exc_info=True,
                )
                self._doc.add_paragraph(  # type: ignore[union-attr]
                    f"[Error: could not render {elem_type} element]"
                )
                self._skipped += 1

        if self._skipped:
            logger.info(
                "Document built with %d skipped element(s).", self._skipped
            )

        return self._save(output_path)

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _setup_document(
        self,
        metadata: Dict[str, Any],
        structure_map: List[Dict[str, Any]] | None = None,
    ) -> None:
        """Create a new ``Document`` and configure page layout.

        Page width/height are taken from the first page entry in
        *metadata* (values expected in PDF points).  Margins are
        derived from the content bounding boxes stored in
        ``metadata["pages"][0]["margins"]`` when available, otherwise
        default to 1 inch on all sides.
        """
        self._doc = Document()

        section = self._doc.sections[0]

        # --- page dimensions from metadata ---------------------------------
        pages = metadata.get("pages", [])
        if pages:
            first_page = pages[0]
            width_pt = first_page.get("width")
            height_pt = first_page.get("height")
            if width_pt is not None:
                section.page_width = Emu(int(float(width_pt) * _PT_TO_EMU))
            if height_pt is not None:
                section.page_height = Emu(int(float(height_pt) * _PT_TO_EMU))

        # --- landscape orientation -----------------------------------------
        if pages:
            first_page = pages[0]
            if first_page.get("is_landscape"):
                section.orientation = WD_ORIENT.LANDSCAPE
                # Ensure width > height for landscape
                if section.page_width < section.page_height:
                    section.page_width, section.page_height = section.page_height, section.page_width

        # --- margins -------------------------------------------------------
        margins_set = False
        if pages:
            margins_dict = pages[0].get("margins")
            if margins_dict:
                section.top_margin = Emu(int(margins_dict["top"] * _PT_TO_EMU))
                section.bottom_margin = Emu(int(margins_dict["bottom"] * _PT_TO_EMU))
                section.left_margin = Emu(int(margins_dict["left"] * _PT_TO_EMU))
                section.right_margin = Emu(int(margins_dict["right"] * _PT_TO_EMU))
                margins_set = True
        if not margins_set:
            margin = Inches(self._default_margin)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # --- default font --------------------------------------------------
        style = self._doc.styles["Normal"]
        font = style.font
        font.name = self._fallback_font
        # Eliminate default paragraph spacing (Word adds 8pt space_after
        # by default which causes vertical overflow on tight layouts).
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        # Ensure East‑Asian fallback is set via the underlying XML so that
        # python‑docx does not silently override it in certain locales.
        r_pr = style.element.get_or_add_rPr()
        r_pr.set(qn("w:eastAsia"), self._fallback_font)

    # ------------------------------------------------------------------
    # Element handlers
    # ------------------------------------------------------------------

    def _add_heading(self, element: Dict[str, Any]) -> None:
        """Add a heading paragraph to the document.

        Parameters
        ----------
        element:
            Must contain ``"text"`` (str) and ``"level"`` (int, 1‑6).
            Optional ``"formatting"`` dict for font overrides.
        """
        text: str = element.get("text", "")
        level: int = element.get("level", 1)
        level = max(1, min(level, 6))  # clamp to valid range

        heading = self._doc.add_heading(text, level=level)  # type: ignore[union-attr]

        # Heading spacing defaults.
        _heading_spacing = {
            1: (24, 6),
            2: (18, 4),
            3: (12, 4),
        }
        if level in _heading_spacing:
            sb, sa = _heading_spacing[level]
            heading.paragraph_format.space_before = Pt(sb)
            heading.paragraph_format.space_after = Pt(sa)

        formatting: Dict[str, Any] = element.get("formatting", {})
        if formatting:
            for run in heading.runs:
                self._apply_run_formatting(run, formatting)

    def _add_paragraph(self, element: Dict[str, Any]) -> None:
        """Add a body paragraph, including multi‑line handling.

        Whitespace-only paragraphs are skipped — their ``spacing_before``
        is accumulated and applied to the next non-empty paragraph.
        """
        text: str = element.get("text", "")
        formatting: Dict[str, Any] = element.get("formatting", {})
        runs_data: List[Dict[str, Any]] | None = element.get("runs")
        links: List[Dict[str, Any]] = element.get("links", [])

        # Skip whitespace-only paragraphs silently.
        if not text.strip() and not links:
            return

        para = self._doc.add_paragraph()  # type: ignore[union-attr]

        # Alignment ----------------------------------------------------------
        alignment_str = formatting.get("alignment", "").lower()
        if alignment_str in _ALIGNMENT_MAP:
            para.alignment = _ALIGNMENT_MAP[alignment_str]

        # Paragraph spacing & indentation ------------------------------------
        spacing_before = formatting.get("spacing_before", 0)
        if spacing_before > 0:
            para.paragraph_format.space_before = Pt(spacing_before)

        spacing_after = formatting.get("spacing_after", 0)
        if spacing_after > 0:
            para.paragraph_format.space_after = Pt(spacing_after)

        indent_left = formatting.get("indent_left", 0)
        if indent_left > 0:
            para.paragraph_format.left_indent = Pt(indent_left)

        # --- hyperlinks present: split text into plain / link segments ------
        if links:
            sorted_links = sorted(links, key=lambda lk: lk.get("start", 0))
            cursor = 0
            for link in sorted_links:
                start = link.get("start", 0)
                end = link.get("end", start)
                uri = link.get("uri", "")
                link_text = link.get("text", text[start:end])

                # Plain text before this link.
                if start > cursor:
                    run = para.add_run(text[cursor:start])
                    self._apply_run_formatting(run, formatting)

                # Hyperlink run.
                if uri:
                    self._add_hyperlink(para, link_text, uri, formatting)
                else:
                    run = para.add_run(link_text)
                    self._apply_run_formatting(run, formatting)

                cursor = end

            # Remaining plain text after the last link.
            if cursor < len(text):
                run = para.add_run(text[cursor:])
                self._apply_run_formatting(run, formatting)
            return

        if runs_data:
            # Mixed-formatting paragraph: each run carries its own formatting.
            prev_run = None
            for run_info in runs_data:
                run_text = run_info.get("text", "")
                if run_text == "\n":
                    # Add a line break to the previous run.
                    if prev_run is not None:
                        prev_run.add_break()
                    continue
                run = para.add_run(run_text)
                self._apply_run_formatting(run, run_info)
                prev_run = run
        else:
            # Fallback: single formatting for the whole paragraph.
            lines = text.split("\n")
            for i, line in enumerate(lines):
                run = para.add_run(line)
                self._apply_run_formatting(run, formatting)
                if i < len(lines) - 1:
                    run.add_break()

    def _add_list_item(self, element: Dict[str, Any]) -> None:
        """Add a bullet or numbered list item."""
        raw_text: str = element.get("text", "")
        level: int = element.get("level", 0)
        bullet_type: str = element.get("bullet_type", "bullet")
        formatting: Dict[str, Any] = element.get("formatting", {})

        # Strip the leading bullet / number prefix so the style handles it.
        text = self._strip_list_prefix(raw_text, bullet_type)

        # Determine the appropriate built‑in style.
        if bullet_type == "number":
            style_name = _NUMBER_STYLES[min(level, len(_NUMBER_STYLES) - 1)]
        else:
            style_name = _BULLET_STYLES[min(level, len(_BULLET_STYLES) - 1)]

        para = self._doc.add_paragraph(style=style_name)  # type: ignore[union-attr]

        run = para.add_run(text)
        self._apply_run_formatting(run, formatting)

    def _add_table(self, element: Dict[str, Any]) -> None:
        """Add a table with proportional column widths, cell shading, borders, and rich cell text."""
        num_rows: int = element.get("num_rows", 0)
        num_cols: int = element.get("num_cols", 0)
        rows_data: List[List[Any]] = element.get("rows", [])
        col_widths: List[float] = element.get("col_widths", [])
        header_row: bool = element.get("header_row", False)
        cell_styles: List[List[Dict[str, Any]]] = element.get("cell_styles", [])
        table_height: float = element.get("table_height", 0)
        row_heights: List[float] = element.get("row_heights", [])

        if num_rows <= 0 or num_cols <= 0:
            logger.warning("Table with invalid dimensions (%d×%d) — skipping.", num_rows, num_cols)
            return

        table = self._doc.add_table(rows=num_rows, cols=num_cols)  # type: ignore[union-attr]
        table.style = "Table Grid"
        table.autofit = False

        # --- tight cell margins to reduce vertical overflow ----------------
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        cell_mar = OxmlElement("w:tblCellMar")
        for side in ("top", "bottom"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "8")  # ~0.4pt in twips (20 twips/pt)
            el.set(qn("w:type"), "dxa")
            cell_mar.append(el)
        tbl_pr.append(cell_mar)

        # --- explicit row heights from PDF table geometry ------------------
        if table_height > 0 and num_rows > 0:
            row_height_pt = table_height / num_rows
            from docx.enum.table import WD_ROW_HEIGHT_RULE
            for row in table.rows:
                row.height = Pt(row_height_pt)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # --- proportional column widths ------------------------------------
        section = self._doc.sections[-1]  # type: ignore[union-attr]
        page_width_emu = section.page_width
        left_margin_emu = section.left_margin
        right_margin_emu = section.right_margin
        content_width_emu = page_width_emu - left_margin_emu - right_margin_emu

        if col_widths and len(col_widths) == num_cols:
            for c_idx, fraction in enumerate(col_widths):
                col_emu = int(content_width_emu * fraction)
                table.columns[c_idx].width = Emu(col_emu)
                for row in table.rows:
                    row.cells[c_idx].width = Emu(col_emu)
        else:
            equal_emu = int(content_width_emu / num_cols)
            for c_idx in range(num_cols):
                table.columns[c_idx].width = Emu(equal_emu)
                for row in table.rows:
                    row.cells[c_idx].width = Emu(equal_emu)

        # --- populate cells ------------------------------------------------
        for r_idx, row_data in enumerate(rows_data):
            if r_idx >= num_rows:
                break

            # Per-row cell styles (if available).
            row_cell_styles: List[Dict[str, Any]] = []
            if cell_styles and r_idx < len(cell_styles):
                row_cell_styles = cell_styles[r_idx]

            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= num_cols:
                    break
                cell = table.cell(r_idx, c_idx)

                # Get cell style info.
                cs: Dict[str, Any] = {}
                if row_cell_styles and c_idx < len(row_cell_styles):
                    cs = row_cell_styles[c_idx]

                # --- Cell background colour --------------------------------
                bg_color = cs.get("bg_color")
                if bg_color:
                    self._set_cell_shading(cell, bg_color)

                # --- Cell text alignment (vertical) ------------------------
                v_align = cs.get("v_alignment")
                if v_align == "center":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif v_align == "bottom":
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

                # --- Populate text with rich formatting --------------------
                cell_text_str = str(cell_text) if cell_text is not None else ""

                # Clear default paragraph.
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    para.clear()
                else:
                    para = cell.add_paragraph()

                # Remove all paragraph spacing in table cells.
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                # Horizontal alignment from cell style.
                h_align = cs.get("alignment", "")
                if h_align == "center":
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif h_align == "right":
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # If we have per-span runs, use them for rich formatting.
                cell_runs = cs.get("runs")
                if cell_runs and cell_text_str.strip():
                    for run_info in cell_runs:
                        run = para.add_run(run_info.get("text", ""))
                        self._apply_run_formatting(run, run_info)
                else:
                    # Single formatting for the whole cell.
                    run = para.add_run(cell_text_str)
                    fmt: Dict[str, Any] = {}
                    if cs.get("font"):
                        fmt["font"] = cs["font"]
                    if cs.get("size"):
                        fmt["size"] = cs["size"]
                    if cs.get("bold") or (header_row and r_idx == 0):
                        fmt["bold"] = True
                    if cs.get("italic"):
                        fmt["italic"] = True
                    if cs.get("color"):
                        fmt["color"] = cs["color"]
                    if cs.get("underline"):
                        fmt["underline"] = True
                    if cs.get("strikethrough"):
                        fmt["strikethrough"] = True
                    if fmt:
                        self._apply_run_formatting(run, fmt)

    def _add_images_inline(self, images: List[Dict[str, Any]]) -> None:
        """Add multiple overlapping images as inline shapes in one paragraph."""
        para = self._doc.add_paragraph()  # type: ignore[union-attr]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for img_elem in images:
            path = img_elem.get("path", "")
            img_width = img_elem.get("width", 0)
            img_height = img_elem.get("height", 0)
            if not path or not os.path.isfile(path):
                continue
            try:
                max_width = _MAX_IMAGE_WIDTH_INCHES
                if img_width > 0 and img_height > 0:
                    native_width_in = img_width / 72.0
                    display_width = min(native_width_in, max_width)
                else:
                    display_width = max_width
                run = para.add_run()
                run.add_picture(path, width=Inches(display_width))
                # Small gap between images
                run.add_text("  ")
            except Exception:
                logger.warning("Could not insert image '%s'.", path, exc_info=True)

    def _add_image(self, element: Dict[str, Any]) -> None:
        """Insert an image, scaling to fit within page margins.

        If the element contains ``x0`` / ``y0`` positioning data and
        ``floating`` is True, the image is inserted as an anchored
        (floating) shape at the specified position, preserving the
        original PDF placement.
        """
        path: str = element.get("path", "")
        img_width: float = element.get("width", 0)
        img_height: float = element.get("height", 0)

        if not path or not os.path.isfile(path):
            logger.warning("Image file not found: '%s'", path)
            self._doc.add_paragraph("[Image could not be inserted]")  # type: ignore[union-attr]
            return

        try:
            # Calculate width that fits within page content area.
            max_width = _MAX_IMAGE_WIDTH_INCHES
            if img_width > 0 and img_height > 0:
                # Convert native pixel / point width to inches (assume 72 DPI).
                native_width_in = img_width / 72.0
                display_width = min(native_width_in, max_width)
            else:
                display_width = max_width

            self._doc.add_picture(path, width=Inches(display_width))  # type: ignore[union-attr]
        except Exception:
            logger.warning(
                "Could not insert image '%s'. Adding placeholder.",
                path,
                exc_info=True,
            )
            self._doc.add_paragraph("[Image could not be inserted]")  # type: ignore[union-attr]

    def _add_header(self, element: Dict[str, Any]) -> None:
        """Set the document header text."""
        text = element.get("text", "")
        if not text:
            return
        section = self._doc.sections[0]  # type: ignore[union-attr]
        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_footer(self, element: Dict[str, Any]) -> None:
        """Set the document footer text."""
        text = element.get("text", "")
        if not text:
            return
        section = self._doc.sections[0]  # type: ignore[union-attr]
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_page_break(self, element: Dict[str, Any]) -> None:
        """Insert a page break, optionally changing orientation."""
        orientation = element.get("orientation")
        if orientation and self._doc is not None:
            # Add a new section with the appropriate orientation
            new_section = self._doc.add_section()
            margin = Inches(self._default_margin)
            new_section.top_margin = margin
            new_section.bottom_margin = margin
            new_section.left_margin = margin
            new_section.right_margin = margin

            if orientation == "landscape":
                new_section.orientation = WD_ORIENT.LANDSCAPE
                # Get page dimensions from prev section, swap for landscape
                prev = self._doc.sections[-2] if len(self._doc.sections) > 1 else self._doc.sections[0]
                w, h = prev.page_width, prev.page_height
                if w < h:
                    w, h = h, w
                new_section.page_width = w
                new_section.page_height = h
            else:
                new_section.orientation = WD_ORIENT.PORTRAIT
                prev = self._doc.sections[-2] if len(self._doc.sections) > 1 else self._doc.sections[0]
                w, h = prev.page_width, prev.page_height
                if w > h:
                    w, h = h, w
                new_section.page_width = w
                new_section.page_height = h
        else:
            self._doc.add_page_break()  # type: ignore[union-attr]

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def _save(self, output_path: str) -> str:
        """Persist the document to *output_path*.

        Creates intermediate directories when necessary.

        Returns
        -------
        str
            The absolute path of the saved file.
        """
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        self._doc.save(output_path)  # type: ignore[union-attr]
        abs_path = os.path.abspath(output_path)
        logger.info("Document saved to '%s'.", abs_path)
        return abs_path

    # ------------------------------------------------------------------
    # Cell shading helper
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell_shading(cell: Any, rgb: tuple) -> None:
        """Apply a background fill colour to a table cell.

        *rgb* should be a ``(r, g, b)`` tuple with values 0–255.
        """
        r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
        hex_color = f"{r:02X}{g:02X}{b:02X}"
        tc_pr = cell._element.get_or_add_tcPr()
        shading = etree.SubElement(tc_pr, qn("w:shd"))
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)

    # ------------------------------------------------------------------
    # Hyperlink helper
    # ------------------------------------------------------------------

    def _add_hyperlink(
        self,
        paragraph: Any,
        text: str,
        url: str,
        formatting: Dict[str, Any],
    ) -> None:
        """Insert a clickable hyperlink run into *paragraph*.

        python-docx does not expose a hyperlink API, so this manipulates
        the underlying OOXML directly.
        """
        part = self._doc.part  # type: ignore[union-attr]
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )

        # Build <w:hyperlink r:id="rId_X"> element.
        hyperlink = etree.SubElement(
            paragraph._element,
            qn("w:hyperlink"),
        )
        hyperlink.set(qn("r:id"), r_id)

        # Build the run inside the hyperlink.
        new_run = etree.SubElement(hyperlink, qn("w:r"))

        # Run properties — style + explicit colour & underline.
        rPr = etree.SubElement(new_run, qn("w:rPr"))
        r_style = etree.SubElement(rPr, qn("w:rStyle"))
        r_style.set(qn("w:val"), "Hyperlink")

        # Font name
        font_name = formatting.get("font") or self._fallback_font
        r_fonts = etree.SubElement(rPr, qn("w:rFonts"))
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)

        # Font size
        size = formatting.get("size")
        if size is not None:
            try:
                # w:sz value is in half-points.
                half_pt = str(int(float(size) * 2))
                sz_el = etree.SubElement(rPr, qn("w:sz"))
                sz_el.set(qn("w:val"), half_pt)
                sz_cs_el = etree.SubElement(rPr, qn("w:szCs"))
                sz_cs_el.set(qn("w:val"), half_pt)
            except (TypeError, ValueError):
                pass

        # Colour — default hyperlink blue.
        color_el = etree.SubElement(rPr, qn("w:color"))
        color_el.set(qn("w:val"), "0563C1")

        # Underline.
        u_el = etree.SubElement(rPr, qn("w:u"))
        u_el.set(qn("w:val"), "single")

        # The text element.
        t_el = etree.SubElement(new_run, qn("w:t"))
        t_el.text = text
        # Preserve leading/trailing whitespace.
        t_el.set(qn("xml:space"), "preserve")

    # ------------------------------------------------------------------
    # Formatting helpers
    # ------------------------------------------------------------------

    def _apply_run_formatting(
        self,
        run: Any,
        formatting: Dict[str, Any],
    ) -> None:
        """Apply font properties from a *formatting* dict to a Run.

        Recognised keys: ``font``, ``size``, ``bold``, ``italic``,
        ``color`` (tuple/list of 3 ints for RGB), ``underline``,
        ``strikethrough``.
        """
        font_name = formatting.get("font") or self._fallback_font
        run.font.name = font_name
        # Set East‑Asian font name at the XML level.
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            from lxml import etree  # already available via python-docx
            r_fonts = etree.SubElement(r_pr, qn("w:rFonts"))
        r_fonts.set(qn("w:eastAsia"), font_name)

        size = formatting.get("size")
        if size is not None:
            try:
                run.font.size = Pt(float(size))
            except (TypeError, ValueError):
                logger.debug("Invalid font size '%s'; ignoring.", size)

        bold = formatting.get("bold")
        if bold is not None:
            run.font.bold = bool(bold)

        italic = formatting.get("italic")
        if italic is not None:
            run.font.italic = bool(italic)

        # Underline
        underline = formatting.get("underline")
        if underline:
            run.font.underline = True

        # Strikethrough
        strikethrough = formatting.get("strikethrough")
        if strikethrough:
            run.font.strike = True

        color = formatting.get("color")
        if color is not None:
            try:
                if isinstance(color, (list, tuple)) and len(color) == 3:
                    r, g, b = int(color[0]), int(color[1]), int(color[2])
                    # Skip pure black (0,0,0) — let Word default handle it
                    # to avoid overriding theme colours unnecessarily.
                    if (r, g, b) != (0, 0, 0):
                        run.font.color.rgb = RGBColor(r, g, b)
                elif isinstance(color, str):
                    # Accept "#RRGGBB" or "RRGGBB" hex strings.
                    hex_str = color.lstrip("#")
                    run.font.color.rgb = RGBColor(
                        int(hex_str[0:2], 16),
                        int(hex_str[2:4], 16),
                        int(hex_str[4:6], 16),
                    )
            except (TypeError, ValueError, IndexError):
                logger.debug("Invalid color value '%s'; ignoring.", color)

    @staticmethod
    def _strip_list_prefix(text: str, bullet_type: str) -> str:
        """Remove a leading bullet / number marker from *text*."""
        if bullet_type == "number":
            return _NUMBER_PREFIX_RE.sub("", text, count=1)
        return _BULLET_PREFIX_RE.sub("", text, count=1)

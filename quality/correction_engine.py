"""Correction engine — applies fixes based on AI-detected differences.

Takes a list of structured differences (from ``AIComparator``) and
modifies the DOCX document to address them.  This is the auto-fix
component of the compare → fix → re-render loop.
"""

import logging
import re
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

logger = logging.getLogger(__name__)

# Constants
MAX_ISSUE_LOG_LENGTH = 50

# Common color name → hex mapping
_COLOR_NAMES = {
    "red": "FF0000",
    "green": "00FF00",
    "blue": "0000FF",
    "black": "000000",
    "white": "FFFFFF",
    "yellow": "FFFF00",
    "orange": "FFA500",
    "purple": "800080",
    "gray": "808080",
    "grey": "808080",
}


def _parse_hex_color(value: str) -> tuple[int, int, int] | None:
    """Parse a color string into (r, g, b) ints, or None on failure.

    Supports: "#FF0000", "FF0000", "red", "rgb(255,0,0)".
    """
    if not value:
        return None
    value = value.strip()

    # Named colors
    lower = value.lower()
    if lower in _COLOR_NAMES:
        h = _COLOR_NAMES[lower]
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    # rgb(r,g,b)
    m = re.match(r"rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", lower)
    if m:
        return int(m.group(1)), int(m.group(2)), int(m.group(3))

    # Hex (with or without #)
    h = value.lstrip("#").strip()
    if re.fullmatch(r"[0-9A-Fa-f]{6}", h):
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    return None


def _parse_pt_value(value: str) -> float | None:
    """Extract a numeric pt value from strings like '14pt', '14', '14.5px'."""
    if not value:
        return None
    m = re.search(r"(\d+(?:\.\d+)?)", value)
    if m:
        return float(m.group(1))
    return None


class CorrectionEngine:
    """Apply corrections to a DOCX based on detected differences.

    Parameters
    ----------
    docx_path : str
        Path to the DOCX file to modify.
    """

    def __init__(self, docx_path: str) -> None:
        self.docx_path = docx_path
        self._doc = Document(docx_path)
        self._fixes_applied = 0

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def apply_fixes(
        self,
        differences: list[dict[str, Any]],
    ) -> int:
        """Apply corrections for the given list of differences.

        Parameters
        ----------
        differences : list[dict]
            Flat list of difference dicts across all pages.
            Each must have ``type``, ``issue``, ``area``, ``severity``.

        Returns
        -------
        int
            Number of fixes successfully applied.
        """
        if not differences:
            logger.debug("No differences to fix.")
            return 0

        self._fixes_applied = 0

        # Validate input
        valid_diffs = []
        for i, diff in enumerate(differences):
            if not isinstance(diff, dict):
                logger.warning("Skipping non-dict difference at index %d", i)
                continue
            if "type" not in diff:
                logger.warning("Skipping difference without 'type' field: %s", diff.get("issue", "unknown"))
                continue
            valid_diffs.append(diff)

        if not valid_diffs:
            logger.warning("No valid differences to fix after validation.")
            return 0

        logger.info("Applying corrections for %d difference(s).", len(valid_diffs))

        # Group by type for batch processing.
        for diff in valid_diffs:
            diff_type = diff.get("type", "")
            try:
                handler = self._handlers.get(diff_type)
                if handler:
                    before = self._fixes_applied
                    handler(self, diff)
                    after = self._fixes_applied
                    if after > before:
                        logger.debug(
                            "Applied fix for '%s': %s",
                            diff_type,
                            diff.get("issue", "")[:MAX_ISSUE_LOG_LENGTH],
                        )
                else:
                    logger.debug(
                        "No handler for difference type '%s': %s",
                        diff_type,
                        diff.get("issue", "")[:MAX_ISSUE_LOG_LENGTH],
                    )
            except Exception as e:
                logger.warning(
                    "Failed to apply fix for '%s': %s (error: %s)",
                    diff_type,
                    diff.get("issue", "")[:MAX_ISSUE_LOG_LENGTH],
                    str(e),
                )

        if self._fixes_applied > 0:
            try:
                self._doc.save(self.docx_path)
                logger.info(
                    "Correction engine: %d fix(es) applied and saved.",
                    self._fixes_applied,
                )
            except Exception as e:
                logger.error("Failed to save corrected document: %s", e)
                raise

        return self._fixes_applied

    # ------------------------------------------------------------------
    # Target-finding helpers
    # ------------------------------------------------------------------

    def _all_paragraphs(self):
        """Yield every paragraph in the document (body + table cells)."""
        yield from self._doc.paragraphs
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def _text_matches(self, text_content: str, paragraph_text: str) -> bool:
        """Fuzzy substring match — shorter string inside longer string."""
        a = text_content.lower()
        b = paragraph_text.lower()
        if not a or not b:
            return False
        # Shorter string should be a substring of the longer one
        if len(a) <= len(b):
            return a in b
        return b in a

    def _find_target_runs(self, diff: dict[str, Any]) -> list:
        """Find runs matching the diff's text_content."""
        text_content = diff.get("text_content", "").strip()

        # If text_content is too short, return all runs (broad match)
        if len(text_content) < 3:
            runs = []
            for para in self._all_paragraphs():
                runs.extend(para.runs)
            return runs

        matched_runs = []
        for para in self._all_paragraphs():
            if self._text_matches(text_content, para.text):
                matched_runs.extend(para.runs)

        if not matched_runs:
            logger.warning(
                "No matching runs found for text_content='%s'",
                text_content[:MAX_ISSUE_LOG_LENGTH],
            )
        return matched_runs

    def _find_target_paragraphs(self, diff: dict[str, Any]) -> list:
        """Find paragraphs matching the diff's text_content."""
        text_content = diff.get("text_content", "").strip()

        # If text_content is too short, return all paragraphs (broad match)
        if len(text_content) < 3:
            return list(self._all_paragraphs())

        matched = []
        for para in self._all_paragraphs():
            if self._text_matches(text_content, para.text):
                matched.append(para)

        if not matched:
            logger.warning(
                "No matching paragraphs found for text_content='%s'",
                text_content[:MAX_ISSUE_LOG_LENGTH],
            )
        return matched

    def _find_target_cells(self, diff: dict[str, Any]) -> list:
        """Find table cells matching the diff's text_content."""
        text_content = diff.get("text_content", "").strip()

        matched = []
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(text_content) < 3:
                        matched.append(cell)
                    elif self._text_matches(text_content, cell.text):
                        matched.append(cell)

        if not matched:
            logger.warning(
                "No matching cells found for text_content='%s'",
                text_content[:MAX_ISSUE_LOG_LENGTH],
            )
        return matched

    # ------------------------------------------------------------------
    # Fix handlers
    # ------------------------------------------------------------------

    def _fix_font_size(self, diff: dict[str, Any]) -> None:
        """Set font size to the expected value from the diff."""
        try:
            expected = diff.get("expected_value", "")
            size_val = _parse_pt_value(expected)
            if size_val is None:
                logger.warning(
                    "Cannot parse font size from expected_value='%s'", expected
                )
                return

            runs = self._find_target_runs(diff)
            for run in runs:
                run.font.size = Pt(size_val)
                self._fixes_applied += 1

            if runs:
                logger.debug(
                    "Set font size to %.1fpt on %d run(s)", size_val, len(runs)
                )
        except Exception as e:
            logger.warning("_fix_font_size failed: %s", e)

    def _fix_bold(self, diff: dict[str, Any]) -> None:
        """Set or remove bold formatting based on expected_value."""
        try:
            expected = diff.get("expected_value", "").lower()
            issue = diff.get("issue", "").lower()

            # Determine target bold state
            if expected:
                should_be_bold = (
                    "bold" in expected
                    and "not bold" not in expected
                    and "no bold" not in expected
                )
            else:
                # Infer from issue text
                should_be_bold = any(
                    kw in issue
                    for kw in ("missing bold", "should be bold", "not bold in converted")
                )

            runs = self._find_target_runs(diff)
            for run in runs:
                run.font.bold = should_be_bold
                self._fixes_applied += 1

            if runs:
                logger.debug(
                    "Set bold=%s on %d run(s)", should_be_bold, len(runs)
                )
        except Exception as e:
            logger.warning("_fix_bold failed: %s", e)

    def _fix_italic(self, diff: dict[str, Any]) -> None:
        """Set or remove italic formatting based on expected_value."""
        try:
            expected = diff.get("expected_value", "").lower()
            issue = diff.get("issue", "").lower()

            if expected:
                should_be_italic = (
                    "italic" in expected
                    and "not italic" not in expected
                    and "no italic" not in expected
                )
            else:
                should_be_italic = any(
                    kw in issue
                    for kw in ("missing italic", "should be italic", "not italic in converted")
                )

            runs = self._find_target_runs(diff)
            for run in runs:
                run.font.italic = should_be_italic
                self._fixes_applied += 1

            if runs:
                logger.debug(
                    "Set italic=%s on %d run(s)", should_be_italic, len(runs)
                )
        except Exception as e:
            logger.warning("_fix_italic failed: %s", e)

    def _fix_alignment(self, diff: dict[str, Any]) -> None:
        """Fix paragraph alignment from expected_value."""
        try:
            expected = diff.get("expected_value", "").lower()
            issue = diff.get("issue", "").lower()

            # Determine target alignment
            source = expected if expected else issue
            _align_map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            target = None
            for keyword, align_val in _align_map.items():
                if keyword in source:
                    target = align_val
                    break

            if target is None:
                logger.warning(
                    "Cannot determine alignment from expected_value='%s' or issue",
                    diff.get("expected_value", ""),
                )
                return

            paragraphs = self._find_target_paragraphs(diff)
            for para in paragraphs:
                para.alignment = target
                self._fixes_applied += 1

            if paragraphs:
                logger.debug(
                    "Set alignment to '%s' on %d paragraph(s)", source, len(paragraphs)
                )
        except Exception as e:
            logger.warning("_fix_alignment failed: %s", e)

    def _fix_spacing(self, diff: dict[str, Any]) -> None:
        """Set paragraph spacing to the expected value."""
        try:
            expected = diff.get("expected_value", "")
            size_val = _parse_pt_value(expected)
            if size_val is None:
                logger.warning(
                    "Cannot parse spacing from expected_value='%s'", expected
                )
                return

            issue = diff.get("issue", "").lower()
            # Determine before vs after — default to space_after
            is_before = any(
                kw in issue for kw in ("before", "space before", "above")
            )

            paragraphs = self._find_target_paragraphs(diff)
            for para in paragraphs:
                pf = para.paragraph_format
                if is_before:
                    pf.space_before = Pt(size_val)
                else:
                    pf.space_after = Pt(size_val)
                self._fixes_applied += 1

            if paragraphs:
                which = "space_before" if is_before else "space_after"
                logger.debug(
                    "Set %s=%.1fpt on %d paragraph(s)",
                    which, size_val, len(paragraphs),
                )
        except Exception as e:
            logger.warning("_fix_spacing failed: %s", e)

    def _fix_font_color(self, diff: dict[str, Any]) -> None:
        """Fix font color from expected_value."""
        try:
            expected = diff.get("expected_value", "")
            rgb = _parse_hex_color(expected)
            if rgb is None:
                logger.warning(
                    "Cannot parse font color from expected_value='%s'", expected
                )
                return

            r, g, b = rgb
            runs = self._find_target_runs(diff)
            for run in runs:
                run.font.color.rgb = RGBColor(r, g, b)
                self._fixes_applied += 1

            if runs:
                logger.debug(
                    "Set font color to #%02X%02X%02X on %d run(s)",
                    r, g, b, len(runs),
                )
        except Exception as e:
            logger.warning("_fix_font_color failed: %s", e)

    def _fix_shading(self, diff: dict[str, Any]) -> None:
        """Fix cell shading from expected_value (hex color)."""
        try:
            expected = diff.get("expected_value", "")
            rgb = _parse_hex_color(expected)
            if rgb is None:
                logger.warning(
                    "Cannot parse shading color from expected_value='%s'", expected
                )
                return

            hex_color = "%02X%02X%02X" % rgb
            cells = self._find_target_cells(diff)
            for cell in cells:
                tc_pr = cell._element.get_or_add_tcPr()
                # Remove existing shading if present
                existing = tc_pr.find(qn("w:shd"))
                if existing is not None:
                    tc_pr.remove(existing)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), hex_color)
                tc_pr.append(shd)
                self._fixes_applied += 1

            if cells:
                logger.debug(
                    "Set shading to #%s on %d cell(s)", hex_color, len(cells)
                )
        except Exception as e:
            logger.warning("_fix_shading failed: %s", e)

    def _fix_border(self, diff: dict[str, Any]) -> None:
        """Add or modify table cell borders based on expected_value."""
        try:
            expected = diff.get("expected_value", "").lower()

            # Determine border style
            no_border = any(
                kw in expected for kw in ("no border", "none", "remove")
            )
            # Parse size if present (e.g. "1px solid black" → sz=4)
            sz_val = "4"  # default thin
            if "thick" in expected:
                sz_val = "12"
            elif "medium" in expected:
                sz_val = "8"

            # Parse color
            border_color = "000000"
            rgb = _parse_hex_color(expected)
            if rgb:
                border_color = "%02X%02X%02X" % rgb

            cells = self._find_target_cells(diff)
            for cell in cells:
                tc_pr = cell._element.get_or_add_tcPr()
                # Remove existing borders
                existing = tc_pr.find(qn("w:tcBorders"))
                if existing is not None:
                    tc_pr.remove(existing)

                if no_border:
                    # Set borders to "none"
                    borders = OxmlElement("w:tcBorders")
                    for side in ("top", "left", "bottom", "right"):
                        el = OxmlElement(f"w:{side}")
                        el.set(qn("w:val"), "none")
                        el.set(qn("w:sz"), "0")
                        el.set(qn("w:space"), "0")
                        el.set(qn("w:color"), "auto")
                        borders.append(el)
                    tc_pr.append(borders)
                else:
                    borders = OxmlElement("w:tcBorders")
                    for side in ("top", "left", "bottom", "right"):
                        el = OxmlElement(f"w:{side}")
                        el.set(qn("w:val"), "single")
                        el.set(qn("w:sz"), sz_val)
                        el.set(qn("w:space"), "0")
                        el.set(qn("w:color"), border_color)
                        borders.append(el)
                    tc_pr.append(borders)
                self._fixes_applied += 1

            if cells:
                style = "none" if no_border else f"single sz={sz_val}"
                logger.debug(
                    "Set borders (%s) on %d cell(s)", style, len(cells)
                )
        except Exception as e:
            logger.warning("_fix_border failed: %s", e)

    def _skip_handler(self, diff: dict[str, Any]) -> None:
        """Log that a fix type is not supported for post-hoc correction."""
        diff_type = diff.get("type", "unknown")
        logger.debug(
            "Skipping '%s' fix — not supported for post-hoc correction",
            diff_type,
        )

    # Map from difference type to handler method.
    _handlers = {
        "font_size": _fix_font_size,
        "font_family": _skip_handler,
        "font_color": _fix_font_color,
        "bold": _fix_bold,
        "italic": _fix_italic,
        "underline": _skip_handler,
        "alignment": _fix_alignment,
        "spacing": _fix_spacing,
        "border": _fix_border,
        "shading": _fix_shading,
        "image": _skip_handler,
        "layout": _skip_handler,
        "missing_content": _skip_handler,
        "extra_content": _skip_handler,
    }

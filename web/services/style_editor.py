"""Prompt-driven DOCX style editor with content-preservation safeguards."""

from __future__ import annotations

import hashlib
import json
import logging
import os
import tempfile
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

_STYLE_PROMPT = """You are a DOCX styling assistant.

Task: convert the user's styling request into JSON style instructions.
IMPORTANT: preserve all document content. Do not rewrite or remove text.
Only return styling/layout changes.

Return ONLY valid JSON using this exact schema:
{
  "theme": {
    "font_name": "string or null",
    "font_size_pt": number or null,
    "line_spacing": number or null,
    "space_after_pt": number or null,
    "accent_color": "#RRGGBB or null",
    "margins_in": {"top": number, "bottom": number, "left": number, "right": number} or null
  },
  "heading": {
    "font_name": "string or null",
    "bold": boolean or null,
    "color": "#RRGGBB or null"
  },
  "table": {
    "header_fill": "#RRGGBB or null",
    "header_font_color": "#RRGGBB or null"
  },
  "summary": "short description"
}
"""


@dataclass
class StyleResult:
    summary: str
    rules: dict[str, Any]
    changed: bool


class StyleEditor:
    """Apply style/layout updates to DOCX while preserving content."""

    def apply_prompt(
        self,
        docx_path: str,
        prompt: str,
        api_key: str | None = None,
        model: str = "gemini-2.0-flash",
    ) -> StyleResult:
        if not os.path.isfile(docx_path):
            raise FileNotFoundError(f"DOCX not found: {docx_path}")
        if not prompt.strip():
            raise ValueError("Prompt cannot be empty")

        rules = self._rules_from_prompt(prompt, api_key=api_key, model=model)

        doc_before = Document(docx_path)
        before_sig = self._content_signature(doc_before)

        changed = self._apply_rules(doc_before, rules)

        if not changed:
            return StyleResult(
                summary=rules.get("summary", "No style changes needed."),
                rules=rules,
                changed=False,
            )

        fd, tmp_path = tempfile.mkstemp(prefix="styled_", suffix=".docx")
        os.close(fd)
        try:
            doc_before.save(tmp_path)
            doc_after = Document(tmp_path)
            after_sig = self._content_signature(doc_after)
            if before_sig != after_sig:
                raise RuntimeError("Content lock violated: style edit changed document text")
            os.replace(tmp_path, docx_path)
        finally:
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass

        return StyleResult(
            summary=rules.get("summary", "Style updated."),
            rules=rules,
            changed=True,
        )

    def _rules_from_prompt(
        self,
        prompt: str,
        api_key: str | None,
        model: str,
    ) -> dict[str, Any]:
        if api_key:
            try:
                from google import genai
                from google.genai import types

                client = genai.Client(api_key=api_key)
                response = client.models.generate_content(
                    model=model,
                    contents=[_STYLE_PROMPT, f"User request: {prompt}"],
                    config=types.GenerateContentConfig(
                        max_output_tokens=800,
                        temperature=0.2,
                    ),
                )
                text = (response.text or "").strip()
                parsed = self._parse_json(text)
                if isinstance(parsed, dict):
                    return self._normalize_rules(parsed)
            except Exception as exc:
                logger.warning("Style AI generation failed, using fallback rules: %s", exc)

        return self._fallback_rules(prompt)

    @staticmethod
    def _parse_json(text: str) -> Any:
        cleaned = text.strip()
        if cleaned.startswith("```"):
            lines = [ln for ln in cleaned.splitlines() if not ln.strip().startswith("```")]
            cleaned = "\n".join(lines)
        return json.loads(cleaned)

    def _fallback_rules(self, prompt: str) -> dict[str, Any]:
        lower = prompt.lower()
        accent = "#1f4e79"
        if "green" in lower:
            accent = "#2e7d32"
        elif "red" in lower:
            accent = "#b71c1c"
        elif "purple" in lower:
            accent = "#5e35b1"

        font_name = "Calibri"
        if "serif" in lower:
            font_name = "Times New Roman"
        elif "modern" in lower or "minimal" in lower:
            font_name = "Aptos"

        strict_spacing = "compact" in lower or "tight" in lower

        return {
            "theme": {
                "font_name": font_name,
                "font_size_pt": 11,
                "line_spacing": 1.15 if not strict_spacing else 1.0,
                "space_after_pt": 8 if not strict_spacing else 4,
                "accent_color": accent,
                "margins_in": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            },
            "heading": {
                "font_name": font_name,
                "bold": True,
                "color": accent,
            },
            "table": {
                "header_fill": accent,
                "header_font_color": "#FFFFFF",
            },
            "summary": "Applied fallback theme styling while preserving all text content.",
        }

    def _normalize_rules(self, rules: dict[str, Any]) -> dict[str, Any]:
        result = {
            "theme": rules.get("theme") or {},
            "heading": rules.get("heading") or {},
            "table": rules.get("table") or {},
            "summary": rules.get("summary") or "Applied AI style updates.",
        }
        return result

    @staticmethod
    def _content_signature(doc: Document) -> str:
        paragraphs = [p.text for p in doc.paragraphs]
        table_cells: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells.append(cell.text)
        payload = "\n".join(paragraphs + table_cells)
        return hashlib.sha256(payload.encode("utf-8")).hexdigest()

    def _apply_rules(self, doc: Document, rules: dict[str, Any]) -> bool:
        changed = False

        theme = rules.get("theme", {})
        heading = rules.get("heading", {})
        table = rules.get("table", {})

        accent_rgb = self._hex_to_rgb(theme.get("accent_color"))
        heading_rgb = self._hex_to_rgb(heading.get("color")) or accent_rgb
        header_fill = table.get("header_fill")
        header_font_rgb = self._hex_to_rgb(table.get("header_font_color"))

        normal_style = doc.styles["Normal"] if "Normal" in [s.name for s in doc.styles] else None
        if normal_style is not None:
            if theme.get("font_name"):
                normal_style.font.name = str(theme["font_name"])
                changed = True
            if theme.get("font_size_pt"):
                normal_style.font.size = Pt(float(theme["font_size_pt"]))
                changed = True

        margins = theme.get("margins_in")
        if isinstance(margins, dict):
            for section in doc.sections:
                for key, attr in (("top", "top_margin"), ("bottom", "bottom_margin"), ("left", "left_margin"), ("right", "right_margin")):
                    if key in margins and margins[key] is not None:
                        setattr(section, attr, Inches(float(margins[key])))
                        changed = True

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            para_fmt = paragraph.paragraph_format
            if theme.get("line_spacing") is not None:
                para_fmt.line_spacing = float(theme["line_spacing"])
                changed = True
            if theme.get("space_after_pt") is not None:
                para_fmt.space_after = Pt(float(theme["space_after_pt"]))
                changed = True

            is_heading = style_name.startswith("Heading")
            for run in paragraph.runs:
                if theme.get("font_name"):
                    run.font.name = str(theme["font_name"])
                    changed = True
                if theme.get("font_size_pt"):
                    run.font.size = Pt(float(theme["font_size_pt"]))
                    changed = True
                if accent_rgb and not is_heading:
                    # Keep body text neutral by default; only apply accent to headings.
                    pass

                if is_heading:
                    if heading.get("font_name"):
                        run.font.name = str(heading["font_name"])
                        changed = True
                    if heading.get("bold") is not None:
                        run.bold = bool(heading["bold"])
                        changed = True
                    if heading_rgb is not None:
                        run.font.color.rgb = heading_rgb
                        changed = True

            if is_heading:
                para_fmt.keep_with_next = True
                para_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                changed = True

        for tbl in doc.tables:
            tbl.autofit = False
            changed = True
            if not tbl.rows:
                continue
            header_row = tbl.rows[0]
            if header_fill:
                for cell in header_row.cells:
                    self._set_cell_shading(cell, header_fill)
                    changed = True
            if header_font_rgb is not None:
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = header_font_rgb
                            run.bold = True
                            changed = True

        return changed

    @staticmethod
    def _hex_to_rgb(value: Any) -> RGBColor | None:
        if not value or not isinstance(value, str):
            return None
        text = value.strip().lstrip("#")
        if len(text) != 6:
            return None
        try:
            return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))
        except ValueError:
            return None

    @staticmethod
    def _set_cell_shading(cell: Any, fill_hex: str) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#"))
        tc_pr.append(shd)

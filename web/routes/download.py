"""Download & preview routes."""

import logging
import os
from flask import Blueprint, send_file, jsonify, abort

from web.services.converter import ConverterService

logger = logging.getLogger(__name__)

download_bp = Blueprint("download", __name__)

_converter: ConverterService | None = None


def init_download(converter: ConverterService) -> None:
    global _converter
    _converter = converter


@download_bp.route("/api/download/<job_id>")
def download(job_id: str):
    """Serve the converted .docx file."""
    assert _converter is not None
    job = _converter.get_job(job_id)
    if job is None:
        return jsonify({"error": "Job not found."}), 404
    if job.status != "complete":
        return jsonify({"error": "Conversion not finished yet."}), 409
    if not os.path.isfile(job.output_path):
        return jsonify({"error": "Output file missing."}), 404

    download_name = os.path.splitext(job.original_filename)[0] + ".docx"
    return send_file(
        job.output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
    )


@download_bp.route("/api/preview/<job_id>/<int:page_num>")
def preview_page(job_id: str, page_num: int):
    """Render a PDF page as a PNG thumbnail for side-by-side preview.

    ``page_num`` is 0-based.
    """
    assert _converter is not None
    job = _converter.get_job(job_id)
    if job is None:
        return jsonify({"error": "Job not found."}), 404

    if not os.path.isfile(job.input_path):
        return jsonify({"error": "Source PDF not found."}), 404

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(job.input_path)
        if page_num < 0 or page_num >= doc.page_count:
            doc.close()
            abort(404)

        page = doc[page_num]
        # Render at 1.5x zoom for readable thumbnails.
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        doc.close()

        return img_bytes, 200, {"Content-Type": "image/png"}

    except Exception as exc:
        logger.exception("Preview render failed: %s", exc)
        return jsonify({"error": str(exc)}), 500


@download_bp.route("/api/docx-preview/<job_id>")
def docx_preview(job_id: str):
    """Return a structured JSON representation of the DOCX for preview.

    This gives the browser enough information to render a side-by-side
    comparison (PDF page image on the left, structured DOCX content on
    the right) without needing to open Word.
    """
    assert _converter is not None
    job = _converter.get_job(job_id)
    if job is None:
        return jsonify({"error": "Job not found."}), 404
    if job.status != "complete":
        return jsonify({"error": "Conversion not finished."}), 409
    if not os.path.isfile(job.output_path):
        return jsonify({"error": "Output file missing."}), 404

    try:
        from docx import Document
        doc = Document(job.output_path)

        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": p.style.name if p.style else "Normal",
                })

        tables = []
        for t in doc.tables:
            rows = []
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(cells)
            tables.append({
                "rows": rows,
                "num_rows": len(t.rows),
                "num_cols": len(t.columns),
            })

        return jsonify({
            "paragraphs": paragraphs,
            "tables": tables,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables),
        })

    except Exception as exc:
        logger.exception("DOCX preview failed: %s", exc)
        return jsonify({"error": str(exc)}), 500

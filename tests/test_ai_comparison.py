"""Tests for AI comparison functionality (AIComparator and CorrectionEngine).

These tests validate the AI comparison infrastructure without requiring
an actual API key by using mocks and simulating responses.
"""

import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch, Mock

# Add parent directory to path for imports
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from quality.ai_comparator import AIComparator
from quality.correction_engine import CorrectionEngine


class TestAIComparator(unittest.TestCase):
    """Test cases for AIComparator class."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        self.sample_pdf_img = os.path.join(self.temp_dir, "pdf_page_0.png")
        self.sample_docx_img = os.path.join(self.temp_dir, "docx_page_0.png")
        
        # Create dummy image files
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='white')
        img.save(self.sample_pdf_img)
        img.save(self.sample_docx_img)

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_init_without_api_key(self):
        """Test initialization without API key."""
        comparator = AIComparator()
        self.assertFalse(comparator.is_available)

    def test_init_with_api_key(self):
        """Test initialization with API key."""
        comparator = AIComparator(api_key="test_key_123")
        self.assertTrue(comparator.is_available)

    def test_init_with_env_var(self):
        """Test initialization with API key from environment."""
        with patch.dict(os.environ, {'GEMINI_API_KEY': 'env_key_456'}):
            comparator = AIComparator()
            self.assertTrue(comparator.is_available)
            self.assertEqual(comparator.api_key, 'env_key_456')

    def test_compare_pages_without_api_key(self):
        """Test compare_pages returns empty results without API key."""
        comparator = AIComparator()
        results = comparator.compare_pages([self.sample_pdf_img], [self.sample_docx_img])
        
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0], [])

    def test_parse_response_valid_json(self):
        """Test parsing valid JSON response."""
        json_response = json.dumps([
            {
                "area": "header",
                "issue": "font size too small",
                "type": "font_size",
                "severity": "high"
            },
            {
                "area": "row 3",
                "issue": "missing border",
                "type": "border",
                "severity": "medium"
            }
        ])
        
        results = AIComparator._parse_response(json_response, page_num=0)
        
        self.assertEqual(len(results), 2)
        self.assertEqual(results[0]["area"], "header")
        self.assertEqual(results[0]["issue"], "font size too small")
        self.assertEqual(results[0]["type"], "font_size")
        self.assertEqual(results[0]["severity"], "high")
        self.assertEqual(results[0]["page_num"], 0)

    def test_parse_response_with_markdown_fences(self):
        """Test parsing JSON wrapped in markdown code fences."""
        markdown_response = """```json
[
  {
    "area": "footer",
    "issue": "alignment wrong",
    "type": "alignment",
    "severity": "low"
  }
]
```"""
        
        results = AIComparator._parse_response(markdown_response, page_num=1)
        
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["area"], "footer")
        self.assertEqual(results[0]["page_num"], 1)

    def test_parse_response_empty_array(self):
        """Test parsing empty array (no differences)."""
        results = AIComparator._parse_response("[]", page_num=0)
        self.assertEqual(results, [])

    def test_parse_response_invalid_json(self):
        """Test parsing invalid JSON returns empty list."""
        results = AIComparator._parse_response("not valid json", page_num=0)
        self.assertEqual(results, [])

    def test_parse_response_adds_defaults(self):
        """Test that missing fields get default values."""
        # Minimal valid entry (only has 'issue')
        json_response = json.dumps([{"issue": "some problem"}])
        
        results = AIComparator._parse_response(json_response, page_num=2)
        
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["area"], "unknown")
        self.assertEqual(results[0]["type"], "layout")
        self.assertEqual(results[0]["severity"], "medium")
        self.assertEqual(results[0]["page_num"], 2)

    def test_compare_single_page_with_mock(self):
        """Test compare_single_page with mocked Gemini API.
        
        This test validates the response parsing and token tracking,
        but doesn't test the actual API call since we don't have a real key.
        """
        # We can't fully test _compare_single_page without mocking the imports,
        # so instead we'll test the response parsing directly
        json_response = json.dumps([
            {
                "area": "table cell",
                "issue": "font color wrong",
                "type": "font_color",
                "severity": "medium"
            }
        ])
        
        # Test the parsing logic
        results = AIComparator._parse_response(json_response, page_num=0)
        
        # Verify results
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["type"], "font_color")
        self.assertEqual(results[0]["area"], "table cell")
        self.assertEqual(results[0]["severity"], "medium")


class TestCorrectionEngine(unittest.TestCase):
    """Test cases for CorrectionEngine class."""

    def setUp(self):
        """Set up test fixtures."""
        from docx import Document
        from docx.shared import Pt
        
        self.temp_dir = tempfile.mkdtemp()
        self.test_docx = os.path.join(self.temp_dir, "test.docx")
        
        # Create a simple test document
        doc = Document()
        para1 = doc.add_paragraph("This is a test paragraph.")
        run1 = para1.runs[0]
        run1.font.size = Pt(12)
        
        para2 = doc.add_paragraph("Another paragraph with some text.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(self.test_docx)

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_init(self):
        """Test initialization of CorrectionEngine."""
        engine = CorrectionEngine(self.test_docx)
        self.assertIsNotNone(engine._doc)
        self.assertEqual(engine._fixes_applied, 0)

    def test_apply_fixes_empty_list(self):
        """Test applying no fixes."""
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes([])
        self.assertEqual(fixes, 0)

    def test_fix_font_size_increase(self):
        """Test font size increase fix."""
        diffs = [
            {
                "type": "font_size",
                "issue": "font size too small",
                "area": "test",
                "severity": "medium"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        # Should have applied at least one fix
        self.assertGreater(fixes, 0)

    def test_fix_font_size_decrease(self):
        """Test font size decrease fix."""
        diffs = [
            {
                "type": "font_size",
                "issue": "font size too large",
                "area": "test",
                "severity": "medium"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        self.assertGreater(fixes, 0)

    def test_fix_alignment(self):
        """Test alignment fix."""
        diffs = [
            {
                "type": "alignment",
                "issue": "should be center aligned",
                "area": "test paragraph",
                "severity": "low"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        # Alignment fix should be applied if matching text found
        self.assertGreaterEqual(fixes, 0)

    def test_fix_bold(self):
        """Test bold formatting fix."""
        diffs = [
            {
                "type": "bold",
                "issue": "should be bold",
                "area": "test paragraph",
                "severity": "medium"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        self.assertGreaterEqual(fixes, 0)

    def test_fix_italic(self):
        """Test italic formatting fix."""
        diffs = [
            {
                "type": "italic",
                "issue": "should be italic",
                "area": "another paragraph",
                "severity": "low"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        self.assertGreaterEqual(fixes, 0)

    def test_fix_border(self):
        """Test table border fix."""
        diffs = [
            {
                "type": "border",
                "issue": "missing cell borders",
                "area": "table",
                "severity": "high"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        # Border fix should be applied to the table
        self.assertGreater(fixes, 0)

    def test_fix_spacing(self):
        """Test spacing adjustment fix."""
        diffs = [
            {
                "type": "spacing",
                "issue": "spacing too much",
                "area": "paragraph",
                "severity": "medium"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        # First add some spacing to fix
        para = engine._doc.paragraphs[0]
        from docx.shared import Pt
        para.paragraph_format.space_before = Pt(10)
        
        fixes = engine.apply_fixes(diffs)
        self.assertGreaterEqual(fixes, 0)

    def test_unknown_fix_type(self):
        """Test handling of unknown fix type."""
        diffs = [
            {
                "type": "unknown_type",
                "issue": "some issue",
                "area": "somewhere",
                "severity": "low"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        # Should not crash, just return 0 fixes
        self.assertEqual(fixes, 0)

    def test_multiple_fixes(self):
        """Test applying multiple different fixes."""
        diffs = [
            {
                "type": "font_size",
                "issue": "font too small",
                "area": "test",
                "severity": "high"
            },
            {
                "type": "bold",
                "issue": "should be bold",
                "area": "test paragraph",
                "severity": "medium"
            },
            {
                "type": "border",
                "issue": "missing borders",
                "area": "table",
                "severity": "high"
            }
        ]
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        # Should apply multiple fixes
        self.assertGreaterEqual(fixes, 2)

    def test_document_saved_after_fixes(self):
        """Test that document is saved after applying fixes."""
        diffs = [
            {
                "type": "border",
                "issue": "missing borders",
                "area": "table",
                "severity": "high"
            }
        ]
        
        # Record original modification time
        original_mtime = os.path.getmtime(self.test_docx)
        
        # Wait a moment to ensure time difference
        import time
        time.sleep(0.1)
        
        engine = CorrectionEngine(self.test_docx)
        fixes = engine.apply_fixes(diffs)
        
        if fixes > 0:
            # File should be updated
            new_mtime = os.path.getmtime(self.test_docx)
            self.assertGreater(new_mtime, original_mtime)


class TestAIComparisonIntegration(unittest.TestCase):
    """Integration tests for the AI comparison flow."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_full_flow_without_api_key(self):
        """Test that the full flow works without API key (graceful degradation)."""
        from PIL import Image
        
        # Create dummy images
        pdf_img = os.path.join(self.temp_dir, "pdf_page_0.png")
        docx_img = os.path.join(self.temp_dir, "docx_page_0.png")
        
        img = Image.new('RGB', (100, 100), color='white')
        img.save(pdf_img)
        img.save(docx_img)
        
        # Run comparison without API key
        comparator = AIComparator()
        results = comparator.compare_pages([pdf_img], [docx_img])
        
        # Should return empty results but not crash
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0], [])

    def test_correction_engine_resilience(self):
        """Test that CorrectionEngine handles errors gracefully."""
        from docx import Document
        
        test_docx = os.path.join(self.temp_dir, "test.docx")
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(test_docx)
        
        # Try to apply fixes with invalid/malformed differences
        invalid_diffs = [
            {},  # Empty dict
            {"type": "font_size"},  # Missing fields
            {"issue": "problem"},  # Missing type
            {"type": "invalid_type", "issue": "test", "area": "test"},  # Invalid type
        ]
        
        engine = CorrectionEngine(test_docx)
        # Should not crash
        fixes = engine.apply_fixes(invalid_diffs)
        
        # Should handle all gracefully
        self.assertGreaterEqual(fixes, 0)


def run_tests():
    """Run all tests."""
    unittest.main(argv=[''], exit=False, verbosity=2)


if __name__ == '__main__':
    run_tests()

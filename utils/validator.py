"""Output validation module.

Validates generated DOCX files for structural integrity and consistency
with the source PDF information.
"""

import logging
import os
import zipfile
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# Standard entries expected in a valid DOCX (Open XML) archive.
REQUIRED_DOCX_ENTRIES = frozenset(
    {
        "[Content_Types].xml",
        "word/document.xml",
    }
)


class OutputValidator:
    """Validates DOCX output files for correctness and consistency."""

    def validate(
        self,
        docx_path: str,
        pdf_info: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Validate a DOCX file and return a detailed report.

        Args:
            docx_path: Path to the DOCX file to validate.
            pdf_info: Optional dictionary returned by
                :pyclass:`~utils.pdf_info.PDFInfo.analyze` for
                cross-referencing metadata (e.g. page count).

        Returns:
            A validation-report dict with the following keys:
                - ``valid`` (bool): Overall validation result.
                - ``file_exists`` (bool): Whether the file exists.
                - ``file_size_kb`` (float): File size in kilobytes.
                - ``is_valid_docx`` (bool): Passes ZIP/XML structure checks.
                - ``issues`` (list[str]): Hard-failure descriptions.
                - ``warnings`` (list[str]): Non-fatal observations.
                - ``summary`` (str): Human-readable one-line summary.
        """
        issues: List[str] = []
        warnings: List[str] = []

        report: Dict[str, Any] = {
            "valid": False,
            "file_exists": False,
            "file_size_kb": 0.0,
            "is_valid_docx": False,
            "issues": issues,
            "warnings": warnings,
            "summary": "",
        }

        # 1. File existence and size ----------------------------------------
        if not os.path.isfile(docx_path):
            issues.append(f"File does not exist: {docx_path}")
            report["summary"] = "Validation failed: file not found."
            return report

        report["file_exists"] = True
        file_size = os.path.getsize(docx_path)
        report["file_size_kb"] = round(file_size / 1024, 2)

        if file_size == 0:
            issues.append("File is empty (0 bytes).")
            report["summary"] = "Validation failed: file is empty."
            return report

        # 2. ZIP structure --------------------------------------------------
        if not zipfile.is_zipfile(docx_path):
            issues.append("File is not a valid ZIP archive.")
            report["summary"] = "Validation failed: not a ZIP archive."
            return report

        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                names = set(zf.namelist())
        except zipfile.BadZipFile as exc:
            issues.append(f"Corrupt ZIP file: {exc}")
            report["summary"] = "Validation failed: corrupt ZIP."
            return report

        # 3. Required DOCX entries ------------------------------------------
        missing = REQUIRED_DOCX_ENTRIES - names
        if missing:
            for entry in sorted(missing):
                issues.append(f"Missing required entry: {entry}")
            report["summary"] = (
                "Validation failed: missing standard DOCX entries."
            )
            return report

        report["is_valid_docx"] = True

        # 4. python-docx opening test --------------------------------------
        try:
            from docx import Document

            doc = Document(docx_path)
            paragraph_count = len(doc.paragraphs)
            logger.debug(
                "DOCX opened successfully: %d paragraphs.", paragraph_count
            )
        except ImportError:
            warnings.append(
                "python-docx is not installed; skipping Document open check."
            )
        except Exception as exc:
            issues.append(f"python-docx could not open file: {exc}")
            report["is_valid_docx"] = False

        # 5. Cross-reference with pdf_info ----------------------------------
        if pdf_info is not None:
            self._cross_check(pdf_info, report, warnings)

        # Final status
        report["valid"] = len(issues) == 0
        report["summary"] = self._build_summary(report)

        log_fn = logger.info if report["valid"] else logger.warning
        log_fn("Validation result for '%s': %s", docx_path, report["summary"])

        return report

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _cross_check(
        pdf_info: Dict[str, Any],
        report: Dict[str, Any],
        warnings: List[str],
    ) -> None:
        """Compare DOCX output against source PDF metadata.

        Args:
            pdf_info: Source PDF analysis dict.
            report: The current validation report (modified in-place).
            warnings: Warning list to append to.
        """
        pdf_pages = pdf_info.get("page_count")
        if pdf_pages is not None and pdf_pages > 0:
            if report["file_size_kb"] < 1:
                warnings.append(
                    f"Output file is very small ({report['file_size_kb']} KB) "
                    f"for a {pdf_pages}-page PDF."
                )

        if pdf_info.get("is_encrypted"):
            warnings.append("Source PDF was encrypted; output may be partial.")

    def quality_score(
        self,
        docx_path: str,
        pdf_info: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Run validation and compute a quality fidelity score.

        Returns a dict with the standard validation report keys plus:
            - ``quality_score`` (int): 0-100 overall quality estimate.
            - ``quality_level`` (str): "Perfect" / "Excellent" / "Good" / "Fair".
            - ``metrics`` (dict): Individual metric scores.
        """
        report = self.validate(docx_path, pdf_info)

        # If basic validation failed, score is 0.
        if not report["valid"]:
            report["quality_score"] = 0
            report["quality_level"] = "Invalid"
            report["metrics"] = {}
            return report

        metrics: Dict[str, Any] = {}
        score = 100  # Start perfect, deduct for issues.

        try:
            from docx import Document

            doc = Document(docx_path)

            # --- Content metrics -------------------------------------------
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            image_count = sum(
                1 for rel in doc.part.rels.values()
                if "image" in rel.reltype
            )

            metrics["paragraphs"] = para_count
            metrics["tables"] = table_count
            metrics["images"] = image_count

            # Check against PDF info if available.
            if pdf_info:
                pdf_pages = pdf_info.get("page_count", 0)
                pdf_has_tables = pdf_info.get("has_tables", False)
                pdf_has_images = pdf_info.get("has_images", False)

                # Deduct if no paragraphs for a multi-page PDF.
                if pdf_pages > 0 and para_count == 0:
                    score -= 30
                    metrics["text_extraction"] = "failed"
                else:
                    metrics["text_extraction"] = "ok"

                # Deduct if PDF has tables but DOCX doesn't.
                if pdf_has_tables and table_count == 0:
                    score -= 15
                    metrics["table_extraction"] = "missing"
                elif pdf_has_tables:
                    metrics["table_extraction"] = "ok"
                else:
                    metrics["table_extraction"] = "n/a"

                # Deduct if PDF has images but DOCX doesn't.
                if pdf_has_images and image_count == 0:
                    score -= 15
                    metrics["image_extraction"] = "missing"
                elif pdf_has_images:
                    metrics["image_extraction"] = "ok"
                else:
                    metrics["image_extraction"] = "n/a"

                # Check paragraph-to-page ratio (expect ≥2 paras per page).
                if pdf_pages > 0:
                    ratio = para_count / pdf_pages
                    metrics["para_per_page"] = round(ratio, 1)
                    if ratio < 1:
                        score -= 10

            # --- Formatting metrics ----------------------------------------
            styled_headings = sum(
                1 for p in doc.paragraphs
                if p.style and p.style.name.startswith("Heading")
            )
            metrics["headings"] = styled_headings
            if styled_headings == 0 and para_count > 10:
                score -= 5  # Likely missing heading detection.

            # Check for runs with explicit formatting.
            formatted_runs = 0
            total_runs = 0
            for p in doc.paragraphs:
                for r in p.runs:
                    total_runs += 1
                    if r.bold or r.italic or r.font.color.rgb:
                        formatted_runs += 1
            metrics["total_runs"] = total_runs
            metrics["formatted_runs"] = formatted_runs

        except ImportError:
            score = max(score - 20, 0)
            metrics["note"] = "python-docx not available for deep inspection"
        except Exception as exc:
            score = max(score - 20, 0)
            metrics["error"] = str(exc)

        # Deduct for warnings.
        score -= min(len(report.get("warnings", [])) * 3, 15)
        score = max(score, 0)

        # Determine quality level.
        if score >= 90:
            level = "Perfect"
        elif score >= 75:
            level = "Excellent"
        elif score >= 60:
            level = "Good"
        else:
            level = "Fair"

        report["quality_score"] = score
        report["quality_level"] = level
        report["metrics"] = metrics
        report["summary"] = f"{level} ({score}/100) — {report['summary']}"

        logger.info(
            "Quality score for '%s': %d/100 (%s)", docx_path, score, level
        )
        return report

    @staticmethod
    def _build_summary(report: Dict[str, Any]) -> str:
        """Build a human-readable one-line summary.

        Args:
            report: The current validation report.

        Returns:
            A summary string.
        """
        issues = report["issues"]
        warnings = report["warnings"]
        size_kb = report["file_size_kb"]

        if report["valid"]:
            parts = [f"Valid DOCX ({size_kb} KB)"]
            if warnings:
                parts.append(f"{len(warnings)} warning(s)")
            return "; ".join(parts) + "."
        else:
            return (
                f"Invalid DOCX: {len(issues)} issue(s), "
                f"{len(warnings)} warning(s)."
            )

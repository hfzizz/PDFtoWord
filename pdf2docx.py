#!/usr/bin/env python3
"""PDF to Word Conversion Tool — Main Orchestrator.

Converts a PDF file to a high-fidelity Word document (.docx), preserving
text formatting, images, tables, headings, lists, and layout as closely
as possible.

Usage
-----
    python pdf2docx.py input.pdf output.docx
    python pdf2docx.py --ocr scanned.pdf output.docx
    python pdf2docx.py --verbose --validate input.pdf output.docx
    python pdf2docx.py --batch C:\\pdfs\\ --output-dir C:\\output\\
    python pdf2docx.py --password "secret" encrypted.pdf output.docx
"""

import argparse
import copy
import glob
import json
import logging
import os
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile
from typing import Any, Callable

import fitz  # PyMuPDF

from extractors.text_extractor import TextExtractor
from extractors.image_extractor import ImageExtractor
from extractors.table_extractor import TableExtractor
from extractors.metadata_extractor import MetadataExtractor
from analyzers.semantic_analyzer import SemanticAnalyzer
from builders.docx_builder import DocxBuilder
from utils.pdf_info import PDFInfo
from utils.ocr_handler import OCRHandler
from utils.progress import ProgressTracker
from utils.validator import OutputValidator
from quality.visual_diff import VisualDiff
from quality.ai_comparator import AIComparator
from quality.correction_engine import CorrectionEngine
from quality.ai_layout_analyzer import AILayoutAnalyzer

logger = logging.getLogger("pdf2docx")

# Path to the default config file shipped alongside this script.
_CONFIG_PATH = os.path.join(os.path.dirname(__file__), "config", "settings.json")


# ------------------------------------------------------------------ #
#  Helpers                                                            #
# ------------------------------------------------------------------ #

def _load_config(config_path: str | None = None) -> dict[str, Any]:
    """Load configuration from a JSON file, falling back to defaults."""
    path = config_path or _CONFIG_PATH
    defaults: dict[str, Any] = {
        "page_size": "auto",
        "ocr_enabled": False,
        "ocr_language": "eng",
        "preserve_fonts": True,
        "image_quality": "original",
        "table_detection": "auto",
        "heading_detection": "both",
        "skip_watermarks": False,
        "verbose": False,
        "fallback_font": "Arial",
        "max_image_dpi": 300,
        "skip_ocr_if_no_tesseract": True,
        "conversion_engine": "custom",
        "ai_comparison": {
            "enabled": False,
            "strategy": "B",
            "model": "gemini-2.0-flash",
            "max_rounds": 3,
        },
        "quality": {
            "mode": "basic",
            "gate": "warn",
            "min_score": 70,
            "use_visual": False,
            "min_visual_ssim": 0.88,
            "engine_fallback": False,
        },
    }
    if os.path.isfile(path):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                user = json.load(fh)
            defaults.update(user)
        except Exception:
            logger.warning("Could not load config from '%s'; using defaults.", path)
    return defaults


def _normalize_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Unpack the ``bbox`` tuple into separate ``x0, y0, x1, y1`` keys.

    The extractors store ``bbox`` as ``(x0, y0, x1, y1)``, but the
    analyzers expect individual keys for sorting and comparison.
    """
    for b in blocks:
        bbox = b.get("bbox")
        if bbox and len(bbox) == 4:
            b["x0"], b["y0"], b["x1"], b["y1"] = bbox
    return blocks


def _filter_text_in_tables(
    text_blocks: list[dict[str, Any]],
    tables: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Remove text blocks whose centre point falls inside a table bbox.

    This prevents table cell text from being duplicated as standalone
    body paragraphs.
    """
    if not tables:
        return text_blocks

    filtered: list[dict[str, Any]] = []
    for block in text_blocks:
        bx0 = block.get("x0", 0)
        by0 = block.get("y0", 0)
        bx1 = block.get("x1", 0)
        by1 = block.get("y1", 0)
        if bx1 == 0 and by1 == 0:
            # Fallback: try bbox tuple.
            bbox = block.get("bbox")
            if bbox and len(bbox) == 4:
                bx0, by0, bx1, by1 = bbox
            else:
                filtered.append(block)
                continue
        cx = (bx0 + bx1) / 2
        cy = (by0 + by1) / 2
        block_page = block.get("page_num", -1)
        in_table = False
        for tbl in tables:
            # Only filter text on the same page as the table.
            if tbl.get("page_num", -2) != block_page:
                continue
            tx0 = tbl.get("x0", 0)
            ty0 = tbl.get("y0", 0)
            tx1 = tbl.get("x1", 0)
            ty1 = tbl.get("y1", 0)
            if tx1 == 0 and ty1 == 0:
                tbl_bbox = tbl.get("bbox")
                if tbl_bbox and len(tbl_bbox) == 4:
                    tx0, ty0, tx1, ty1 = tbl_bbox
                else:
                    continue
            if tx0 <= cx <= tx1 and ty0 <= cy <= ty1:
                in_table = True
                break
        if not in_table:
            filtered.append(block)
    return filtered


def _update_metadata_margins(
    metadata: dict[str, Any],
    text_blocks: list[dict[str, Any]],
    images: list[dict[str, Any]],
    tables: list[dict[str, Any]],
) -> None:
    """Compute per-page content margins and store in *metadata*.

    Scans all extracted items for their bounding boxes to determine
    how close the content gets to the page edges.  The computed margins
    are stored under ``metadata["pages"][n]["margins"]`` as a dict with
    ``top``, ``bottom``, ``left``, ``right`` in PDF points.  A minimum
    of 14 pt (~0.2 in) is enforced.
    """
    _MIN_MARGIN = 14.0
    pages = metadata.get("pages", [])
    if not pages:
        return

    # Collect bounding boxes per page.
    per_page: dict[int, list[tuple[float, float, float, float]]] = {}

    for block in text_blocks:
        pn = block.get("page_num", 0)
        x0, y0, x1, y1 = _get_bbox(block)
        if x1 > x0 and y1 > y0:
            per_page.setdefault(pn, []).append((x0, y0, x1, y1))

    for img in images:
        pn = img.get("page_num", 0)
        x0, y0, x1, y1 = _get_bbox(img)
        if x1 > x0 and y1 > y0:
            per_page.setdefault(pn, []).append((x0, y0, x1, y1))

    for tbl in tables:
        pn = tbl.get("page_num", 0)
        x0, y0, x1, y1 = _get_bbox(tbl)
        if x1 > x0 and y1 > y0:
            per_page.setdefault(pn, []).append((x0, y0, x1, y1))

    for pn, bboxes in per_page.items():
        if pn >= len(pages):
            continue
        pg = pages[pn]
        pw = float(pg.get("width", 612))
        ph = float(pg.get("height", 792))

        min_x = min(b[0] for b in bboxes)
        min_y = min(b[1] for b in bboxes)
        max_x = max(b[2] for b in bboxes)
        max_y = max(b[3] for b in bboxes)

        pg["margins"] = {
            "top": max(_MIN_MARGIN, min_y),
            "bottom": max(_MIN_MARGIN, ph - max_y),
            "left": max(_MIN_MARGIN, min_x),
            "right": max(_MIN_MARGIN, pw - max_x),
        }


def _get_bbox(item: dict[str, Any]) -> tuple[float, float, float, float]:
    """Extract bounding box from an extracted item dict."""
    bbox = item.get("bbox")
    if bbox and len(bbox) >= 4:
        return float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
    return (
        float(item.get("x0", 0)),
        float(item.get("y0", 0)),
        float(item.get("x1", 0)),
        float(item.get("y1", 0)),
    )


def _convert_with_pdf2docx_library(
    input_path: str,
    output_path: str,
    password: str | None = None,
    progress_callback: Callable[[str, int, int, str], None] | None = None,
) -> str:
    """Convert using the external ``pdf2docx`` library CLI in current env."""
    exe_name = "pdf2docx.exe" if os.name == "nt" else "pdf2docx"
    interpreter_dir = Path(sys.executable).resolve().parent
    candidate = interpreter_dir / exe_name
    cli_path = str(candidate) if candidate.is_file() else shutil.which("pdf2docx")
    if not cli_path:
        raise RuntimeError(
            "pdf2docx library CLI not found in current environment. "
            "Install package with: pip install pdf2docx"
        )

    if progress_callback is not None:
        progress_callback("building", 0, 1, "Converting with pdf2docx library…")

    cmd = [cli_path, "convert", input_path, output_path]
    if password:
        cmd.extend(["--password", password])

    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        check=False,
        cwd=tempfile.gettempdir(),
    )
    if proc.returncode != 0:
        stderr = (proc.stderr or "").strip()
        stdout = (proc.stdout or "").strip()
        details = stderr or stdout or "Unknown error"
        raise RuntimeError(f"pdf2docx library conversion failed: {details}")

    if not os.path.isfile(output_path):
        raise RuntimeError("pdf2docx library finished without producing output DOCX")

    if progress_callback is not None:
        progress_callback("building", 1, 1, "pdf2docx library conversion complete")

    return output_path


def _run_docx_postprocess(docx_path: str, mode: str) -> dict[str, int]:
    """Run deterministic post-processing fixes on a DOCX file."""
    if mode == "off":
        return {
            "blank_paragraphs_removed": 0,
            "spacing_clamped": 0,
            "heading_keep_with_next": 0,
            "table_autofit_disabled": 0,
        }

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        logger.warning("python-docx unavailable; skipping deterministic postprocess.")
        return {
            "blank_paragraphs_removed": 0,
            "spacing_clamped": 0,
            "heading_keep_with_next": 0,
            "table_autofit_disabled": 0,
        }

    doc = Document(docx_path)
    stats = {
        "blank_paragraphs_removed": 0,
        "spacing_clamped": 0,
        "heading_keep_with_next": 0,
        "table_autofit_disabled": 0,
    }

    # 1) Remove excess consecutive empty paragraphs.
    empty_run = 0
    for paragraph in list(doc.paragraphs):
        text = (paragraph.text or "").strip()
        if text:
            empty_run = 0
            continue
        empty_run += 1
        if empty_run > 1:
            element = paragraph._element
            element.getparent().remove(element)
            stats["blank_paragraphs_removed"] += 1

    # 2) Clamp extreme paragraph spacing values.
    for paragraph in doc.paragraphs:
        para_fmt = paragraph.paragraph_format
        for attr in ("space_before", "space_after"):
            val = getattr(para_fmt, attr)
            if val is not None and getattr(val, "pt", 0) > 24:
                setattr(para_fmt, attr, Pt(12))
                stats["spacing_clamped"] += 1

        if mode == "strict" and para_fmt.line_spacing is not None:
            line_val = para_fmt.line_spacing
            try:
                numeric = float(line_val)
            except (TypeError, ValueError):
                numeric = None
            if numeric is not None and numeric > 2.0:
                para_fmt.line_spacing = 1.5
                stats["spacing_clamped"] += 1

    # 3) Keep headings with next paragraph to reduce page-break drift.
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            stats["heading_keep_with_next"] += 1

    # 4) Stabilize table layout.
    for table in doc.tables:
        table.autofit = False
        stats["table_autofit_disabled"] += 1

    doc.save(docx_path)
    return stats


def _composite_quality_score(
    base_quality_score: int,
    visual_ssim: float | None,
    use_visual: bool,
) -> int:
    """Build a deterministic composite quality score (0-100)."""
    if not use_visual or visual_ssim is None:
        return int(base_quality_score)
    visual_score = max(0.0, min(1.0, visual_ssim)) * 100.0
    # Weight structural quality higher than visual rendering score.
    blended = (base_quality_score * 0.8) + (visual_score * 0.2)
    return int(round(max(0.0, min(100.0, blended))))


# ------------------------------------------------------------------ #
#  Core pipeline                                                      #
# ------------------------------------------------------------------ #

def convert_pdf(
    input_path: str,
    output_path: str,
    config: dict[str, Any],
    password: str | None = None,
    validate: bool = False,
    visual_validate: bool = False,
    ai_enhance: bool = False,
    ai_strategy: str = "B",
    ai_compare: bool | None = None,
    conversion_engine: str = "custom",
    progress_callback: Callable[[str, int, int, str], None] | None = None,
) -> str:
    """Run the full conversion pipeline for a single PDF.

    Parameters
    ----------
    input_path : str
        Path to the source PDF.
    output_path : str
        Destination path for the ``.docx`` file.
    config : dict
        Application configuration.
    password : str | None
        Optional password for encrypted PDFs.
    validate : bool
        If ``True``, run validation on the output file.
    visual_validate : bool
        If ``True``, render both documents and compute SSIM scores.
    ai_enhance : bool
        If ``True``, enable AI-powered enhancement steps (requires
        ``GEMINI_API_KEY``).
    ai_strategy : str
        ``"A"`` for post-build correction loop,
        ``"B"`` for pre-build AI-guided layout analysis (default).
    ai_compare : bool | None
        Deprecated alias for ``ai_enhance`` (kept for compatibility).
    conversion_engine : str
        ``"custom"`` for this project's pipeline, ``"pdf2docx_lib"`` to use
        the external ``pdf2docx`` library converter.
    progress_callback : callable | None
        Optional ``(stage, current, total, message) -> None`` callback
        invoked at each pipeline stage so callers (e.g. Web UI) can
        report progress.

    Returns
    -------
    str
        The absolute path of the generated ``.docx``.
    """
    def _progress(stage: str, current: int, total: int, message: str) -> None:
        if progress_callback is not None:
            progress_callback(stage, current, total, message)

    # Auto-flushing print so output appears immediately in web server threads.
    def _print(*args: Any, **kwargs: Any) -> None:
        kwargs.setdefault("flush", True)
        print(*args, **kwargs)

    input_path = os.path.abspath(input_path)
    output_path = os.path.abspath(output_path)
    logger.info("Converting '%s' → '%s'", input_path, output_path)

    if ai_compare is not None:
        ai_enhance = ai_enhance or ai_compare
    conversion_engine = (conversion_engine or "custom").lower()
    if conversion_engine not in {"custom", "pdf2docx_lib"}:
        logger.warning("Unknown conversion_engine '%s'; falling back to custom.",
                       conversion_engine)
        conversion_engine = "custom"

    quality_cfg = config.get("quality", {}) or {}
    quality_mode = str(quality_cfg.get("mode", "basic")).lower()
    quality_gate = str(quality_cfg.get("gate", "warn")).lower()
    quality_min_score = int(quality_cfg.get("min_score", 70))
    quality_use_visual = bool(quality_cfg.get("use_visual", False))
    quality_min_visual_ssim = float(quality_cfg.get("min_visual_ssim", 0.88))
    quality_engine_fallback = bool(quality_cfg.get("engine_fallback", False))

    if quality_mode not in {"off", "basic", "strict"}:
        quality_mode = "basic"
    if quality_gate not in {"off", "warn", "fail"}:
        quality_gate = "warn"

    if quality_use_visual and not visual_validate:
        visual_validate = True

    if not os.path.isfile(input_path):
        logger.error("Input file not found: '%s'", input_path)
        sys.exit(1)

    # ── Stage 0: Quick analysis ──────────────────────────────────────
    pdf_info_analyzer = PDFInfo()
    info = pdf_info_analyzer.analyze(input_path)
    logger.info(
        "PDF info — pages: %d, text-based: %s, has images: %s, has tables: %s",
        info.get("page_count", 0),
        info.get("is_text_based"),
        info.get("has_images"),
        info.get("has_tables"),
    )

    # ── Open PDF document ────────────────────────────────────────────
    try:
        doc = fitz.open(input_path)
    except Exception as exc:
        logger.error("Cannot open PDF '%s': %s", input_path, exc)
        sys.exit(1)

    if doc.is_encrypted:
        if password:
            if not doc.authenticate(password):
                logger.error("Incorrect password for '%s'.", input_path)
                doc.close()
                sys.exit(1)
        else:
            logger.error("PDF is encrypted. Supply password with --password.")
            doc.close()
            sys.exit(1)

    page_count = doc.page_count
    progress = ProgressTracker(total=page_count, description="Extracting pages")
    temp_dir: str | None = None
    saved_path: str | None = None

    # Fast path: use external pdf2docx library converter.
    if conversion_engine == "pdf2docx_lib":
        if ai_enhance and ai_strategy.upper() == "B":
            logger.warning(
                "Strategy B requires custom pipeline analysis; using library engine "
                "without Strategy B pre-build overrides."
            )
        _progress("building", page_count, page_count,
                  "Building Word document with pdf2docx library…")
        progress.close()
        doc.close()
        saved_path = _convert_with_pdf2docx_library(
            input_path,
            output_path,
            password=password,
            progress_callback=lambda stage, current, total, message: _progress(
                stage, current, total, message
            ),
        )

    if conversion_engine == "custom":
        # ── Stage 1: Extract metadata ────────────────────────────────────
        _progress("analyzing", 0, page_count, "Extracting metadata…")
        meta_extractor = MetadataExtractor()
        metadata = meta_extractor.extract(doc)

        # ── Stage 2: Content extraction (per-page) ───────────────────────
        text_extractor = TextExtractor()
        image_extractor = ImageExtractor(doc)
        table_extractor = TableExtractor()
        ocr_handler = OCRHandler()

        all_text_blocks: list[dict[str, Any]] = []
        all_images: list[dict[str, Any]] = []
        all_tables: list[dict[str, Any]] = []
        all_links: list[dict[str, Any]] = []

        # Create a temporary directory for extracted images.
        temp_dir = tempfile.mkdtemp(prefix="pdf2docx_imgs_")

        ocr_enabled = config.get("ocr_enabled", False)
        ocr_language = config.get("ocr_language", "eng")

        for page_num in range(page_count):
            page = doc[page_num]
            progress.set_description(f"Page {page_num + 1}/{page_count}")

            # ── Text ──
            _progress("extracting", page_num + 1, page_count,
                      f"Extracting page {page_num + 1}/{page_count}")
            text_blocks = text_extractor.extract(page, page_num)

            # If page has barely any text and OCR is enabled, try OCR.
            text_len = sum(len(b.get("text", "")) for b in text_blocks)
            if text_len < 10 and ocr_enabled:
                if ocr_handler.is_available():
                    logger.info("Page %d: sparse text, running OCR…", page_num + 1)
                    ocr_blocks = ocr_handler.ocr_page(page, language=ocr_language)
                    text_blocks.extend(ocr_blocks)
                else:
                    logger.warning(
                        "Page %d: needs OCR but Tesseract is not installed.", page_num + 1
                    )

            # ── Images ──
            images = image_extractor.extract(page, page_num, temp_dir)
            all_images.extend(images)

            # ── Tables ──
            tables = table_extractor.extract(page, page_num)

            # ── Hyperlinks ──
            links = text_extractor.extract_links(page, page_num)
            all_links.extend(links)

            # ── Filter out text that falls inside table regions ──
            text_blocks = _filter_text_in_tables(text_blocks, tables)

            all_text_blocks.extend(text_blocks)
            all_tables.extend(tables)

            progress.update()

        progress.close()
        doc.close()

        # ── Normalize extracted data ─────────────────────────────────────
        all_text_blocks = _normalize_blocks(all_text_blocks)
        all_images = _normalize_blocks(all_images)
        all_tables = _normalize_blocks(all_tables)

        logger.info(
            "Extraction complete — %d text spans, %d images, %d tables, %d links",
            len(all_text_blocks),
            len(all_images),
            len(all_tables),
            len(all_links),
        )

        # ── Compute content-based margins ────────────────────────────────
        _update_metadata_margins(metadata, all_text_blocks, all_images, all_tables)

        # ── Stage 3: Semantic analysis ───────────────────────────────────
        _progress("analyzing", page_count, page_count, "Semantic analysis…")
        logger.info("Running semantic analysis…")
        analyzer = SemanticAnalyzer(config)
        extracted_data: dict[str, Any] = {
            "text_blocks": all_text_blocks,
            "images": all_images,
            "tables": all_tables,
            "links": all_links,
            "metadata": metadata,
        }
        structure_map = analyzer.analyze(extracted_data)
        logger.info("Structure map: %d elements", len(structure_map))

        # ── Stage 4: Build Word document ─────────────────────────────────
        formatting_overrides: dict[str, Any] = {}

        # Strategy B: AI-guided pre-build analysis
        if ai_enhance and ai_strategy.upper() == "B":
            ai_config = config.get("ai_comparison", {})
            api_key = ai_config.get("api_key") or os.environ.get("GEMINI_API_KEY", "")
            model = ai_config.get("model", "gemini-2.0-flash")

            if api_key:
                _progress("ai_layout", page_count, page_count,
                          "AI layout analysis (Strategy B)…")
                logger.info("Running AI layout analysis (Strategy B)…")
                try:
                    layout_analyzer = AILayoutAnalyzer(
                        api_key=api_key, model=model,
                    )
                    formatting_overrides = layout_analyzer.analyze_layout(
                        input_path, structure_map,
                        progress_callback=lambda cur, total, msg: _progress(
                            "ai_layout", cur, total, msg
                        ),
                    )
                    n_text = len(formatting_overrides.get("text_overrides", {}))
                    n_table = len(formatting_overrides.get("table_overrides", []))
                    logger.info(
                        "AI layout analysis complete — %d text overrides, "
                        "%d table overrides",
                        n_text, n_table,
                    )
                    _print(f"\n  AI layout: {n_text} text overrides, "
                           f"{n_table} table overrides")
                except Exception as exc:
                    logger.warning("AI layout analysis failed: %s — "
                                   "proceeding without overrides.", exc)
            else:
                logger.warning("Strategy B selected but no API key — "
                               "building without AI overrides.")

        _progress("building", page_count, page_count, "Building Word document…")
        logger.info("Building Word document…")
        builder = DocxBuilder(config)
        saved_path = builder.build(structure_map, metadata, output_path,
                                   formatting_overrides=formatting_overrides or None)
        logger.info("Word document saved → '%s'", saved_path)

    # ── Stage 5: Deterministic post-processing + validation ──────────
    quality_report: dict[str, Any] | None = None
    composite_score: int | None = None
    overall_ssim: float | None = None

    if quality_mode != "off":
        _progress("postprocess", page_count, page_count,
                  "Applying deterministic post-processing…")
        post_stats = _run_docx_postprocess(saved_path, quality_mode)
        logger.info(
            "Postprocess stats — blank_removed=%d spacing_clamped=%d "
            "heading_keep_with_next=%d table_autofit_disabled=%d",
            post_stats.get("blank_paragraphs_removed", 0),
            post_stats.get("spacing_clamped", 0),
            post_stats.get("heading_keep_with_next", 0),
            post_stats.get("table_autofit_disabled", 0),
        )

    if validate or quality_mode != "off":
        _progress("validating", page_count, page_count, "Validating output…")
        logger.info("Validating output…")
        validator = OutputValidator()
        quality_report = validator.quality_score(saved_path, info)
        score = quality_report.get("quality_score", "?")
        level = quality_report.get("quality_level", "?")
        if validate or quality_mode != "off":
            _print(f"\n  Quality: {level} ({score}/100)")
        metrics = quality_report.get("metrics", {})
        if metrics:
            logger.info(
                "Metrics: %d paragraphs, %d tables, %d images, %d headings",
                metrics.get("paragraphs", 0),
                metrics.get("tables", 0),
                metrics.get("images", 0),
                metrics.get("headings", 0),
            )
        if not quality_report["valid"]:
            logger.warning("Validation issues:")
            for issue in quality_report.get("issues", []):
                logger.warning("  • %s", issue)
        for warning in quality_report.get("warnings", []):
            logger.warning("  ⚠ %s", warning)

    # ── Stage 6-9: Visual diff + AI comparison (optional) ─────────────
    if visual_validate or ai_enhance:
        _progress("visual_diff", page_count, page_count, "Running visual diff…")
        vd_config = config.get("visual_diff", {})
        dpi = vd_config.get("dpi", 150)
        lo_path = vd_config.get("libreoffice_path", "auto")

        vdiff = VisualDiff(dpi=dpi, libreoffice_path=lo_path)

        # Output dir for diff images alongside the DOCX.
        diff_dir = os.path.splitext(saved_path)[0] + "_visual_diff"
        vd_result = vdiff.compare(input_path, saved_path, diff_dir)

        overall_ssim = vd_result.get("overall_score", 0)
        vd_level = vd_result.get("quality_level", "red")
        page_scores = vd_result.get("page_scores", [])
        pdf_pages = vd_result.get("pdf_page_count", 0)
        docx_pages = vd_result.get("docx_page_count", 0)

        _print(f"\n  Visual SSIM: {overall_ssim:.1%} ({vd_level})")
        if pdf_pages != docx_pages:
            _print(f"  Page count: PDF={pdf_pages}, DOCX={docx_pages}"
                  f" (overflow: {docx_pages - pdf_pages} extra pages)")
        for i, s in enumerate(page_scores):
            tag = "[OK]" if s >= 0.95 else "[!!]" if s < 0.85 else "[--]"
            _print(f"    Page {i + 1}: {s:.1%} {tag}")

        # AI comparison + correction loop (Strategy A only).
        # Strategy B already applied overrides pre-build; just report SSIM.
        if ai_enhance and overall_ssim < 0.95 and ai_strategy.upper() == "A":
            ai_config = config.get("ai_comparison", {})
            api_key = ai_config.get("api_key") or os.environ.get("GEMINI_API_KEY", "")
            model = ai_config.get("model", "gemini-2.0-flash")
            max_rounds = ai_config.get("max_rounds", 3)

            if api_key:
                comparator = AIComparator(
                    api_key=api_key,
                    model=model,
                )

                for round_num in range(1, max_rounds + 1):
                    _progress(
                        "ai_compare", round_num, max_rounds,
                        f"AI comparison round {round_num}/{max_rounds}…",
                    )
                    _print(f"\n  AI comparison round {round_num}/{max_rounds}…")

                    all_diffs = comparator.compare_pages(
                        vd_result["pdf_images"],
                        vd_result["docx_images"],
                    )

                    # Flatten diffs.
                    flat_diffs = [d for page_diffs in all_diffs for d in page_diffs]

                    if not flat_diffs:
                        _print("    No differences detected by AI.")
                        break

                    _print(f"    {len(flat_diffs)} difference(s) found.")
                    for d in flat_diffs[:5]:
                        _print(f"      [{d.get('severity', '?')}] "
                               f"{d.get('area', '?')}: {d.get('issue', '?')}")
                    if len(flat_diffs) > 5:
                        _print(f"      … and {len(flat_diffs) - 5} more")

                    # Apply corrections.
                    _progress(
                        "auto_fix", round_num, max_rounds,
                        f"Applying fixes (round {round_num})…",
                    )
                    engine = CorrectionEngine(saved_path)
                    fixes = engine.apply_fixes(flat_diffs)
                    _print(f"    Applied {fixes} fix(es).")

                    if fixes == 0:
                        _print("    No actionable fixes — stopping loop.")
                        break

                    # Re-render and re-score.
                    vd_result = vdiff.compare(input_path, saved_path, diff_dir)
                    overall_ssim = vd_result.get("overall_score", 0)
                    vd_level = vd_result.get("quality_level", "red")
                    _print(f"    Updated SSIM: {overall_ssim:.1%} ({vd_level})")

                    if overall_ssim >= 0.95:
                        _print("    Target SSIM reached!")
                        break
            else:
                _print("  [SKIP] No GEMINI_API_KEY set — skipping AI comparison.")

    # ── Stage 10: Quality gate + optional deterministic fallback ─────
    if quality_report is not None:
        composite_score = _composite_quality_score(
            int(quality_report.get("quality_score", 0)),
            overall_ssim,
            quality_use_visual,
        )
        logger.info("Composite quality score: %d/100", composite_score)

        gate_failed = composite_score < quality_min_score
        if quality_use_visual and overall_ssim is not None:
            gate_failed = gate_failed or (overall_ssim < quality_min_visual_ssim)

        if gate_failed and quality_engine_fallback:
            alt_engine = "pdf2docx_lib" if conversion_engine == "custom" else "custom"
            logger.warning(
                "Quality gate failed (%d < %d). Trying fallback engine '%s'.",
                composite_score,
                quality_min_score,
                alt_engine,
            )
            fallback_config = copy.deepcopy(config)
            fallback_quality = fallback_config.setdefault("quality", {})
            fallback_quality["engine_fallback"] = False
            fallback_quality["gate"] = "off"

            fd, fallback_path = tempfile.mkstemp(prefix="pdf2docx_fallback_",
                                                 suffix=".docx")
            os.close(fd)
            try:
                convert_pdf(
                    input_path,
                    fallback_path,
                    fallback_config,
                    password=password,
                    validate=False,
                    visual_validate=False,
                    ai_enhance=False,
                    conversion_engine=alt_engine,
                    progress_callback=None,
                )
                validator = OutputValidator()
                fallback_report = validator.quality_score(fallback_path, info)
                fallback_score = int(fallback_report.get("quality_score", 0))
                if fallback_score > composite_score:
                    logger.info(
                        "Fallback engine '%s' improved quality from %d to %d.",
                        alt_engine,
                        composite_score,
                        fallback_score,
                    )
                    shutil.move(fallback_path, saved_path)
                    quality_report = fallback_report
                    composite_score = fallback_score
                    gate_failed = composite_score < quality_min_score
                else:
                    logger.info(
                        "Fallback engine '%s' did not improve quality (%d <= %d).",
                        alt_engine,
                        fallback_score,
                        composite_score,
                    )
            except Exception as exc:
                logger.warning("Fallback conversion failed: %s", exc)
            finally:
                try:
                    if os.path.exists(fallback_path):
                        os.remove(fallback_path)
                except OSError:
                    pass

        if gate_failed and quality_gate == "warn":
            logger.warning(
                "Quality gate warning: composite=%d (min=%d)%s",
                composite_score,
                quality_min_score,
                (
                    f", visual={overall_ssim:.1%} (min={quality_min_visual_ssim:.1%})"
                    if (quality_use_visual and overall_ssim is not None)
                    else ""
                ),
            )
        elif gate_failed and quality_gate == "fail":
            raise RuntimeError(
                f"Quality gate failed: composite={composite_score} < "
                f"min_score={quality_min_score}"
            )

    # ── Clean up temp images ─────────────────────────────────────────
    if temp_dir:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass

    _progress("complete", page_count, page_count, "Conversion complete")
    _print(f"\n[OK] Conversion complete: {saved_path}")
    return saved_path


# ------------------------------------------------------------------ #
#  Batch conversion                                                   #
# ------------------------------------------------------------------ #

def convert_batch(
    input_dir: str,
    output_dir: str,
    config: dict[str, Any],
    password: str | None = None,
    validate: bool = False,
) -> None:
    """Convert all PDF files in *input_dir* to ``.docx``."""
    pattern = os.path.join(input_dir, "*.pdf")
    pdf_files = sorted(glob.glob(pattern))
    if not pdf_files:
        logger.error("No PDF files found in '%s'.", input_dir)
        sys.exit(1)

    os.makedirs(output_dir, exist_ok=True)
    total = len(pdf_files)
    print(f"Batch converting {total} PDF(s)…")

    for i, pdf_path in enumerate(pdf_files, 1):
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        out_path = os.path.join(output_dir, f"{base}.docx")
        print(f"\n[{i}/{total}] {os.path.basename(pdf_path)}")
        try:
            convert_pdf(
                pdf_path,
                out_path,
                config,
                password=password,
                validate=validate,
                conversion_engine=config.get("conversion_engine", "custom"),
            )
        except SystemExit:
            logger.warning("Skipping '%s' due to error.", pdf_path)
        except Exception:
            logger.exception("Unexpected error converting '%s'.", pdf_path)

    print(f"\nBatch complete — {total} file(s) processed → {output_dir}")


# ------------------------------------------------------------------ #
#  CLI entry point                                                    #
# ------------------------------------------------------------------ #

def main() -> None:
    """Parse arguments and run the conversion."""
    parser = argparse.ArgumentParser(
        prog="pdf2docx",
        description="Convert PDF documents to high-fidelity Word (.docx) files.",
    )

    # Positional arguments (optional when --batch is used).
    parser.add_argument("input", nargs="?", help="Input PDF file path.")
    parser.add_argument("output", nargs="?", help="Output .docx file path.")

    # Optional flags.
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Enable OCR for scanned / image-based pages (requires Tesseract).",
    )
    parser.add_argument(
        "--password",
        type=str,
        default=None,
        help="Password for encrypted PDFs.",
    )
    parser.add_argument(
        "--config",
        type=str,
        default=None,
        help="Path to a custom configuration JSON file.",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose (DEBUG) logging.",
    )
    parser.add_argument(
        "--validate",
        action="store_true",
        help="Run validation checks on the output document.",
    )
    parser.add_argument(
        "--batch",
        type=str,
        default=None,
        help="Directory containing PDF files for batch conversion.",
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default=None,
        help="Output directory for batch conversion results.",
    )
    parser.add_argument(
        "--skip-watermarks",
        action="store_true",
        help="Attempt to skip watermark / background images.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing output files without prompting.",
    )
    parser.add_argument(
        "--visual-validate",
        action="store_true",
        help="Render PDF and DOCX, then compute SSIM visual similarity scores.",
    )
    parser.add_argument(
        "--use-pdf2docx-lib",
        action="store_true",
        help="Use external pdf2docx library engine instead of custom pipeline.",
    )
    parser.add_argument(
        "--quality-mode",
        type=str,
        choices=["off", "basic", "strict"],
        default=None,
        help="Deterministic quality mode: off, basic, or strict.",
    )
    parser.add_argument(
        "--quality-gate",
        type=str,
        choices=["off", "warn", "fail"],
        default=None,
        help="Quality gate behavior when score is below threshold.",
    )
    parser.add_argument(
        "--min-quality-score",
        type=int,
        default=None,
        help="Minimum composite quality score (0-100) for quality gate.",
    )
    parser.add_argument(
        "--quality-use-visual",
        action="store_true",
        help="Include visual SSIM in composite quality score.",
    )
    parser.add_argument(
        "--quality-engine-fallback",
        action="store_true",
        help="Retry with alternate engine when quality gate fails.",
    )
    parser.add_argument(
        "--ai-enhance", "--ai-compare",
        dest="ai_enhance",
        action="store_true",
        help="Enable AI-powered enhancement using Gemini (requires GEMINI_API_KEY). "
             "(`--ai-compare` is kept as a compatibility alias)",
    )
    parser.add_argument(
        "--ai-strategy",
        type=str,
        choices=["A", "B"],
        default="B",
        help="AI strategy: A = post-build correction loop, B = pre-build AI-guided layout (default: B).",
    )

    args = parser.parse_args()

    # ── Logging ──────────────────────────────────────────────────────
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
        stream=sys.stderr,
    )

    # ── Config ───────────────────────────────────────────────────────
    config = _load_config(args.config)
    if args.ocr:
        config["ocr_enabled"] = True
    if args.skip_watermarks:
        config["skip_watermarks"] = True
    if args.verbose:
        config["verbose"] = True
    if args.use_pdf2docx_lib:
        config["conversion_engine"] = "pdf2docx_lib"
    quality_config = config.setdefault("quality", {})
    if args.quality_mode is not None:
        quality_config["mode"] = args.quality_mode
    if args.quality_gate is not None:
        quality_config["gate"] = args.quality_gate
    if args.min_quality_score is not None:
        quality_config["min_score"] = max(0, min(100, args.min_quality_score))
    if args.quality_use_visual:
        quality_config["use_visual"] = True
    if args.quality_engine_fallback:
        quality_config["engine_fallback"] = True

    # ── Batch mode ───────────────────────────────────────────────────
    if args.batch:
        out_dir = args.output_dir or os.path.join(args.batch, "docx_output")
        convert_batch(args.batch, out_dir, config, args.password, args.validate)
        return

    # ── Single-file mode ─────────────────────────────────────────────
    if not args.input:
        parser.error("Please provide an input PDF file (or use --batch for batch mode).")

    input_path = args.input
    output_path = args.output
    if not output_path:
        # Default: same name with .docx extension, in the same directory.
        base = os.path.splitext(input_path)[0]
        output_path = base + ".docx"

    # Check overwrite.
    if os.path.exists(output_path) and not args.force:
        resp = input(f"Output file '{output_path}' exists. Overwrite? [y/N] ").strip().lower()
        if resp not in ("y", "yes"):
            print("Aborted.")
            sys.exit(0)

    convert_pdf(
        input_path,
        output_path,
        config,
        password=args.password,
        validate=args.validate,
        visual_validate=args.visual_validate,
        ai_enhance=args.ai_enhance,
        ai_strategy=args.ai_strategy,
        conversion_engine=config.get("conversion_engine", "custom"),
    )


if __name__ == "__main__":
    main()
